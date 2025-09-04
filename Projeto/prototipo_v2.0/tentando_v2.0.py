import tkinter as tk
from tkinter import messagebox, END, ttk, filedialog, Toplevel, Label
import os
import tempfile
from docx2pdf import convert
from docx import Document
from docx.shared import Inches, Pt
from bisect import bisect_right
import sys
import subprocess
from docx.shared import RGBColor
import time
import unicodedata
import re
import sqlite3
import pandas as pd
import brazilcep
from datetime import datetime
import locale
import sv_ttk
import darkdetect

root = tk.Tk()

class AutocompleteEntry(ttk.Entry):
    def __init__(self, autocompleteList, *args, **kwargs):
        self.style = ttk.Style()
        self.style.configure("Big.TEntry", font=("Helvetica", 20))
        # Listbox length
        if 'listboxLength' in kwargs:
            self.listboxLength = kwargs['listboxLength']
            del kwargs['listboxLength']
        else:
            self.listboxLength = 8

        # Custom matches function
        if 'matchesFunction' in kwargs:
            self.matchesFunction = kwargs['matchesFunction']
            del kwargs['matchesFunction']
        else:
            def matches(fieldValue, acListEntry):
                pattern = re.compile('.*' + re.escape(fieldValue) + '.*', re.IGNORECASE)
                return re.match(pattern, acListEntry)

            self.matchesFunction = matches

        ttk.Entry.__init__(self, *args, **kwargs)
        self.focus()

        self.autocompleteList = autocompleteList

        self.var = self["textvariable"]
        if self.var == '':
            self.var = self["textvariable"] = tk.StringVar()

        self.var.trace('w', self.changed)
        self.bind("<Right>", self.selection)
        self.bind("<Up>", self.moveUp)
        self.bind("<Down>", self.moveDown)

        self.listboxUp = False
    def changed(self, name, index, mode):
        if self.var.get() == '':
            if self.listboxUp:
                self.listbox.destroy()
                self.listboxUp = False
        else:
            words = self.comparison()
            if words:
                if not self.listboxUp:
                    self.listbox = tk.Listbox(self.master, width=self["width"], height=self.listboxLength)
                    self.listbox.bind("<Button-1>", self.selection)
                    self.listbox.bind("<Right>", self.selection)
                    self.listbox.place(x=self.winfo_x(), y=self.winfo_y() + self.winfo_height())
                    self.listboxUp = True

                self.listbox.delete(0, END)
                for w in words:
                    self.listbox.insert(END, w)
            else:
                if self.listboxUp:
                    self.listbox.destroy()
                    self.listboxUp = False
    def selection(self, event):
        if self.listboxUp:
            self.var.set(self.listbox.get(tk.ACTIVE))
            self.listbox.destroy()
            self.listboxUp = False
            self.icursor(END)
    def moveUp(self, event):
        if self.listboxUp:
            if self.listbox.curselection() == ():
                index = '0'
            else:
                index = self.listbox.curselection()[0]

            if index != '0':
                self.listbox.selection_clear(first=index)
                index = str(int(index) - 1)

                self.listbox.see(index)  # Scroll!
                self.listbox.selection_set(first=index)
                self.listbox.activate(index)
    def moveDown(self, event):
        if self.listboxUp:
            if self.listbox.curselection() == ():
                index = '0'
            else:
                index = self.listbox.curselection()[0]

            if index != END:
                self.listbox.selection_clear(first=index)
                index = str(int(index) + 1)

                self.listbox.see(index)  # Scroll!
                self.listbox.selection_set(first=index)
                self.listbox.activate(index)
    def comparison(self):
        return [w for w in self.autocompleteList if self.matchesFunction(self.var.get(), w)]

class Funcs():
    def setup_enter_bindings(self, frame):
        # Crie uma lista com todos os widgets na ordem que você quer
        # Substitua pelas variáveis das suas Entrys
        if(frame == "frame_cliente"):
            self.entry_list = [
                self.nome_entry,
                self.nacionalidade_combo,
                self.estado_civil_combo,
                self.profissao_entry,
                self.rg_entry,
                self.cpf_entry,
                self.cep_entry,
                self.n_rua_entry,
                self.telefone_entry,
                self.email_entry,
                self.nome_reu_entry,
                self.cnpj_reu_entry
            ]
        elif(frame == "frame_documento"):
            self.entry_list = [
                self.nome_doc_entry,
                self.tipo_combo,
                self.caminho_doc_entry
            ]
        
        for i in range(len(self.entry_list) - 1):
            current_entry = self.entry_list[i]
            next_entry = self.entry_list[i + 1]
            
            # Cria uma função lambda para passar o próximo widget para a próxima função
            # Isso garante que cada entry saiba para onde ir
            current_entry.bind("<Return>", lambda event, next_widget=next_entry: next_widget.focus_set())
    
    def OnDoubleClick_Cli(self, event):
        self.limpa_cliente()
        self.listaCli.selection()

        for n in self.listaCli.selection():
            col1,col2,col3,col4,col5,col6,col7,col8,col9,col10,col11,col12,col13,col14,col15,col16,col17 = self.listaCli.item(n, 'values')
            self.id_entry.insert(END, col1)
            self.nome_entry.insert(END, col2)
            self.nacionalidade_combo.insert(END, col3)
            self.estado_civil_combo.insert(END, col4)
            self.profissao_entry.insert(END, col5)
            self.rg_entry.insert(END, col6)
            self.cpf_entry.insert(END, col7)
            self.cep_entry.insert(END, col8)
            self.uf_entry.insert(END, col9)
            self.cidade_entry.insert(END, col10)
            self.rua_entry.insert(END, col11)
            self.n_rua_entry.insert(END, col12)
            self.bairro_entry.insert(END, col13)
            self.telefone_entry.insert(END, col14)
            self.email_entry.insert(END, col15)
            self.nome_reu_entry.insert(END, col16)
            self.cnpj_reu_entry.insert(END, col17)
    
    def OnDoubleClick_Doc(self, event):
        self.limpa_documento()
        self.listaDoc.selection()

        for n in self.listaDoc.selection():
            col1,col2,col3 = self.listaDoc.item(n, 'values')
            self.id_doc_entry.insert(END, col1)
            self.nome_doc_entry.insert(END, col2)
            self.tipo_combo.insert(END, col3)
    
    def matches(self, fieldValue, acListEntry):
        pattern = re.compile(re.escape(fieldValue) + '.*', re.IGNORECASE)
        return re.match(pattern, acListEntry)
    
    def limpa_cliente(self):
        self.id_entry.delete(0, END)
        self.nome_entry.delete(0, END)
        self.nacionalidade_combo.delete(0, END)
        self.estado_civil_combo.delete(0, END)
        self.profissao_entry.delete(0, END)
        self.rg_entry.delete(0, END)
        self.cpf_entry.delete(0, END)
        self.cep_entry.delete(0, END)
        self.uf_entry.delete(0, END)
        self.cidade_entry.delete(0, END)
        self.rua_entry.delete(0, END)
        self.n_rua_entry.delete(0, END)
        self.bairro_entry.delete(0, END)
        self.telefone_entry.delete(0, END)
        self.email_entry.delete(0, END)
        self.nome_reu_entry.delete(0, END)
        self.cnpj_reu_entry.delete(0, END)

    def limpa_documento(self):
        self.id_doc_entry.delete(0, END)
        self.nome_doc_entry.delete(0, END)
        self.tipo_combo.delete(0, END)
        self.caminho_doc_entry.delete(0, END)
    
    def conecta_bd(self):
        self.conn = sqlite3.connect("bd_advogados.bd")
        self.cursor =self.conn.cursor()

    def desconecta_bd(self):
        self.conn.close()

    def montaTabelas(self):
        self.conecta_bd()

        # Cria a tabela de clientes
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS clientes (
                cli_id INTEGER PRIMARY KEY,
                cli_nome VARCHAR(100) NOT NULL,
                cli_nacionalidade VARCHAR(50) NOT NULL,
                cli_estado_civil VARCHAR(50) CHECK (cli_estado_civil IN ('solteiro', 'solteira', 'casado', 'casada', 'divorciado', 'divorciada', 'viúvo', 'viúva', 'separado', 'separada', 'união estável', 'separação judicial')),
                cli_profissao VARCHAR(50) NOT NULL,
                cli_rg VARCHAR(10) NOT NULL,
                cli_cpf CHAR(11) UNIQUE NOT NULL,
                cli_cep CHAR(8) NOT NULL,
                cli_uf CHAR(2) NOT NULL,
                cli_cidade VARCHAR(100) NOT NULL,
                cli_logradouro VARCHAR(200) NOT NULL,
                cli_n_rua VARCHAR(10) NOT NULL,
                cli_bairro VARCHAR(100) NOT NULL,
                cli_telefone VARCHAR(10),
                cli_email VARCHAR(100),
                cli_nome_reu VARCHAR(100),
                cli_cnpj_reu VARCHAR(14)
            );
        """)

        # Cria a tabela de documentos
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS documentos (
                doc_id INTEGER PRIMARY KEY,
                doc_nome VARCHAR(100) UNIQUE NOT NULL,
                doc_tipo VARCHAR(20) NOT NULL,
                doc_arquivo BLOB UNIQUE NOT NULL
            );
        """)
                            
        # Cria a tabela de documentos gerados
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS documento_gerado (
                dg_id INTEGER PRIMARY KEY,
                fk_clientes_id INTEGER NOT NULL,
                fk_documentos_id INTEGER NOT NULL,
                dg_nome VARCHAR(100) UNIQUE NOT NULL,
                dg_data_criacao DATE NOT NULL,
                dg_arquivo BLOB UNIQUE NOT NULL,
                FOREIGN KEY (fk_clientes_id) REFERENCES clientes (cli_id),
                FOREIGN KEY (fk_documentos_id) REFERENCES documentos (doc_id)
            );
        """)

        self.conn.commit()
        self.desconecta_bd()
    
    def cepCorreios(self, *args):
        # Funções auxiliares. Elas podem ser movidas para fora da função principal.
        def limpar_campos():
            self.uf_entry.delete(0, END)
            self.cidade_entry.delete(0, END)
            self.rua_entry.delete(0, END)
            self.bairro_entry.delete(0, END)
            self.n_rua_entry.delete(0, END)

        def habilitar_campos():
            self.uf_entry.state(['!disabled'])
            self.cidade_entry.state(['!disabled'])
            self.rua_entry.state(['!disabled'])
            self.bairro_entry.state(['!disabled'])

        def desabilitar_campos():
            self.uf_entry.state(['disabled'])
            self.cidade_entry.state(['disabled'])
            self.rua_entry.state(['disabled'])
            self.bairro_entry.state(['disabled'])

        # Obtém o CEP e remove a formatação
        zipcode = re.sub(r'\D', '', self.cep_entry.get())

        # Se o campo estiver vazio, limpa e habilita os outros campos.
        if len(zipcode) == 0:
            self.cep_entry.state(['!invalid'])
            habilitar_campos()
            limpar_campos()
            return

        # Se o CEP não tiver o tamanho esperado, marca como inválido.
        if len(zipcode) != 8:
            self.cep_entry.state(['invalid'])
            habilitar_campos()
            limpar_campos()
            return
        
        # Tenta buscar o CEP se o tamanho for 8.
        try:
            dadosCep = brazilcep.get_address_from_cep(zipcode)
            
            # Habilita os campos para poderem ser preenchidos e limpa antes de inserir
            habilitar_campos()
            limpar_campos()

            # Insere os dados
            self.uf_entry.insert(END, dadosCep['uf'])
            self.cidade_entry.insert(END, dadosCep['city'])
            self.rua_entry.insert(END, dadosCep['street'])
            self.bairro_entry.insert(END, dadosCep['district'])
            
            if dadosCep['complement'] != '':
                self.n_rua_entry.insert(END, dadosCep['complement'])
            
            # Desabilita os campos após o preenchimento bem-sucedido
            desabilitar_campos()

            # Marca a entry do CEP como válida
            self.cep_entry.state(['!invalid'])

        except KeyError:
            # Se o CEP for inválido
            self.cep_entry.state(['invalid'])
            habilitar_campos()
            limpar_campos()
        except:
            # Erro geral de conexão
            self.cep_entry.state(['invalid'])
            habilitar_campos()
            limpar_campos()
    
    def valida_cpf_cnpj(self, *args):
        entry = self.cnpj_reu_entry # Adapte o nome da sua entry
        documento = entry.get()
        
        # Remove a máscara (pontos, traço, barra)
        documento_limpo = re.sub(r'\D', '', documento)

        # Se o campo estiver vazio, não mostra a borda vermelha
        if not documento_limpo:
            entry.state(['!invalid'])
            return
            
        # Lógica de validação com base no tamanho
        if len(documento_limpo) == 11:
            # A validação é feita como CPF
            if documento_limpo == documento_limpo[0] * 11:
                entry.state(['invalid'])
                return
            
            # Validação do primeiro dígito do CPF
            soma = sum(int(documento_limpo[i]) * (10 - i) for i in range(9))
            resto = soma % 11
            digito_verificador_1 = 11 - resto if resto >= 2 else 0
            if digito_verificador_1 != int(documento_limpo[9]):
                entry.state(['invalid'])
                return

            # Validação do segundo dígito do CPF
            soma = sum(int(documento_limpo[i]) * (11 - i) for i in range(10))
            resto = soma % 11
            digito_verificador_2 = 11 - resto if resto >= 2 else 0
            if digito_verificador_2 != int(documento_limpo[10]):
                entry.state(['invalid'])
                return
            
            # Se for um CPF válido
            entry.state(['!invalid'])

        elif len(documento_limpo) == 14:
            # A validação é feita como CNPJ
            if documento_limpo == documento_limpo[0] * 14:
                entry.state(['invalid'])
                return
            
            # Validação do primeiro dígito do CNPJ
            soma = 0
            multiplicador = 5
            for i in range(12):
                soma += int(documento_limpo[i]) * multiplicador
                multiplicador -= 1
                if multiplicador < 2:
                    multiplicador = 9
            resto = soma % 11
            digito_verificador_1 = 0 if resto < 2 else 11 - resto
            if int(documento_limpo[12]) != digito_verificador_1:
                entry.state(['invalid'])
                return

            # Validação do segundo dígito do CNPJ
            soma = 0
            multiplicador = 6
            for i in range(13):
                soma += int(documento_limpo[i]) * multiplicador
                multiplicador -= 1
                if multiplicador < 2:
                    multiplicador = 9
            resto = soma % 11
            digito_verificador_2 = 0 if resto < 2 else 11 - resto
            if int(documento_limpo[13]) != digito_verificador_2:
                entry.state(['invalid'])
                return

            # Se for um CNPJ válido
            entry.state(['!invalid'])

        else:
            # Tamanho inválido para CPF ou CNPJ
            entry.state(['invalid'])
    
    def valida_email(self, *args):
        email_entry = self.email_entry
        email = email_entry.get().strip()

        # Se o campo estiver vazio, não mostra a borda vermelha
        if not email:
            email_entry.state(['!invalid'])
            return
        
        # Expressão regular para validação de e-mail
        email_regex = re.compile(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$')
        
        if email_regex.fullmatch(email):
            # O e-mail é válido
            email_entry.state(['!invalid'])
        else:
            # O e-mail é inválido
            email_entry.state(['invalid'])
    
    def valida_rg(self, *args):
        rg_entry = self.rg_entry
        rg = rg_entry.get()
        
        # Se o campo estiver vazio, não mostra a borda vermelha
        if not rg:
            rg_entry.state(['!invalid'])
            return
        
        # Remove caracteres não numéricos.
        rg_limpo = re.sub(r'\D', '', rg)
        
        # RG no Brasil tem 9 dígitos (incluindo o dígito verificador)
        if len(rg_limpo) == 9:
            rg_entry.state(['!invalid'])
        else:
            rg_entry.state(['invalid'])
    
    def valida_cpf(self, *args):
        cpf_entry = self.cpf_entry
        cpf = cpf_entry.get()

        # Se o campo estiver vazio, não mostra a borda vermelha
        if not cpf:
            cpf_entry.state(['!invalid'])
            return

        # Remove pontos, traço e outros caracteres não numéricos.
        cpf_limpo = re.sub(r'\D', '', cpf)

        # A validação só acontece quando o campo tiver o tamanho esperado
        if len(cpf_limpo) == 11:
            # Evita CPFs com todos os dígitos iguais, que são inválidos
            if cpf_limpo == cpf_limpo[0] * 11:
                cpf_entry.state(['invalid'])
                return

            # Validação do primeiro dígito verificador
            soma = sum(int(cpf_limpo[i]) * (10 - i) for i in range(9))
            resto = soma % 11
            digito_verificador_1 = 11 - resto if resto >= 2 else 0

            if digito_verificador_1 != int(cpf_limpo[9]):
                cpf_entry.state(['invalid'])
                return

            # Validação do segundo dígito verificador
            soma = sum(int(cpf_limpo[i]) * (11 - i) for i in range(10))
            resto = soma % 11
            digito_verificador_2 = 11 - resto if resto >= 2 else 0

            if digito_verificador_2 != int(cpf_limpo[10]):
                cpf_entry.state(['invalid'])
                return
            
            # Se passar em todas as checagens
            cpf_entry.state(['!invalid'])

        # Se o tamanho não for 11, o campo fica em estado inválido
        else:
            cpf_entry.state(['invalid'])
    
    def variaveis_cli(self):
        def formatar_nome_proprio(texto):
            excecoes = ["da", "de", "do", "dos", "das", "e", "em", "para"]
            palavras = texto.lower().split()
            texto_formatado = []
            
            for palavra in palavras:
                if palavra in excecoes:
                    texto_formatado.append(palavra)
                else:
                    texto_formatado.append(palavra.capitalize())
                    
            return " ".join(texto_formatado)
        
        # Coletando dados dos campos
        self.id = re.sub(r'\D', '', self.id_entry.get().strip())
        self.nome_completo = formatar_nome_proprio(self.nome_entry.get().strip())
        self.nacionalidade = self.nacionalidade_combo.get().strip().lower()
        self.estado_civil = self.estado_civil_combo.get().strip().lower()
        self.profissao = self.profissao_entry.get().strip().lower()
        self.rg = re.sub(r'\D', '', self.rg_entry.get().strip())
        self.cpf = re.sub(r'\D', '', self.cpf_entry.get().strip())
        self.cep = re.sub(r'\D', '', self.cep_entry.get().strip())
        self.uf = self.uf_entry.get().strip().upper()
        self.cidade = formatar_nome_proprio(self.cidade_entry.get().strip())
        self.logradouro = formatar_nome_proprio(self.rua_entry.get().strip())
        self.n_rua = self.n_rua_entry.get().strip()
        self.bairro = formatar_nome_proprio(self.bairro_entry.get().strip())
        self.telefone = re.sub(r'\D', '', self.telefone_entry.get().strip())
        self.email = self.email_entry.get().strip().lower()
        self.nome_reu = formatar_nome_proprio(self.nome_reu_entry.get().strip())
        self.cnpj_reu = re.sub(r'\D', '', self.cnpj_reu_entry.get().strip())

    def upt_cliente(self):
        id_para_atualizar = self.id_entry.get().strip()

        try:
            if not id_para_atualizar:
                messagebox.showwarning("Erro de ID", "Por favor, digite o ID do cliente para atualizar.")
                return
            
            if not id_para_atualizar.isdigit():
                messagebox.showerror("Erro de ID", "O ID deve ser um número inteiro.")
                return

            campos_com_validacao = [
                self.rg_entry, self.cpf_entry, self.cep_entry, self.telefone_entry,
                self.email_entry, self.cnpj_reu_entry
            ]
            
            for campo in campos_com_validacao:
                if 'invalid' in campo.state():
                    messagebox.showwarning("Erro de Validação", "Por favor, preencha corretamente os campos marcados em vermelho.")
                    return

            self.conecta_bd()
            self.cursor.execute("SELECT * FROM clientes WHERE cli_id = ?", (id_para_atualizar,))
            dados_atuais = self.cursor.fetchone()

            if not dados_atuais:
                messagebox.showerror("Erro", f"Cliente com o ID: {id_para_atualizar}não encontrado no banco de dados.")
                return

            self.variaveis_cli()

            nome_completo = self.nome_completo if self.nome_completo else dados_atuais[1]
            nacionalidade = self.nacionalidade if self.nacionalidade else dados_atuais[2]
            estado_civil = self.estado_civil if self.estado_civil else dados_atuais[3]
            profissao = self.profissao if self.profissao else dados_atuais[4]
            rg = self.rg if self.rg else dados_atuais[5]
            cpf = self.cpf if self.cpf else dados_atuais[6]
            cep = self.cep if self.cep else dados_atuais[7]
            uf = self.uf if self.uf else dados_atuais[8]
            cidade = self.cidade if self.cidade else dados_atuais[9]
            logradouro = self.logradouro if self.logradouro else dados_atuais[10]
            n_rua = self.n_rua if self.n_rua else dados_atuais[11]
            bairro = self.bairro if self.bairro else dados_atuais[12]
            telefone = self.telefone if self.telefone else dados_atuais[13]
            email = self.email if self.email else dados_atuais[14]
            nome_reu = self.nome_reu if self.nome_reu else dados_atuais[15]
            cnpj_reu = self.cnpj_reu if self.cnpj_reu else dados_atuais[16]
            
            # --- Bloco de Confirmação ---
            nome_cliente_existente = dados_atuais[1] # Pega o nome do cliente que já existe no banco
            
            confirmacao = messagebox.askyesno("Confirmar Atualização", 
                                            f"Tem certeza que deseja atualizar o cliente?\n"
                                            f"\nNome: {nome_cliente_existente}\nID: {id_para_atualizar}")
            
            if not confirmacao:
                return
            # --- Fim do Bloco de Confirmação ---
            
            self.cursor.execute("""
                UPDATE clientes SET 
                    cli_nome = ?, cli_nacionalidade = ?, cli_estado_civil = ?, cli_profissao = ?, cli_rg = ?, 
                    cli_cpf = ?, cli_cep = ?, cli_uf = ?, cli_cidade = ?, cli_logradouro = ?, cli_n_rua = ?, 
                    cli_bairro = ?, cli_telefone = ?, cli_email = ?, cli_nome_reu = ?, cli_cnpj_reu = ? 
                WHERE cli_id = ?
            """, (
                nome_completo, nacionalidade, estado_civil, profissao, rg, 
                cpf, cep, uf, cidade, logradouro, n_rua, 
                bairro, telefone, email, nome_reu, cnpj_reu, 
                id_para_atualizar
            ))
            
            self.conn.commit()
            messagebox.showinfo("Sucesso", "Cliente atualizado com sucesso!")
            self.select_listaClientes(self.listaCli)
            self.limpa_cliente()
            
        except sqlite3.IntegrityError as e:
            messagebox.showerror("Erro de Integridade", "Erro: O CPF ou CNPJ informado já existe no banco de dados.")
        except Exception as e:
            messagebox.showerror("Erro na Atualização", f"Não foi possível atualizar o cliente.\n\nErro: {e}")
            
        finally:
            self.desconecta_bd()
    
    def del_cliente(self, event):
        id_para_deletar = None
        
        # 1. Tenta obter o ID do campo id_entry (se estiver preenchido)
        id_da_entry = self.id_entry.get().strip()
        
        # 2. Tenta obter o ID da linha selecionada na Treeview
        selecao_treeview = self.listaCli.selection()
        
        try:
            if id_da_entry and selecao_treeview:
                # Caso 1: Ambos estão preenchidos, verifica se são iguais
                id_da_linha = self.listaCli.item(selecao_treeview[0], 'values')[0]
                if str(id_da_entry) != str(id_da_linha):
                    messagebox.showwarning("Aviso de Conflito", 
                                        "O ID no campo (digitado) é diferente do ID da linha selecionada.\n"
                                        "Por favor, use apenas um dos métodos para exclusão.")
                    return
                else:
                    id_para_deletar = id_da_entry

            elif id_da_entry:
                # Caso 2: Somente o campo id_entry está preenchido
                if not id_da_entry.isdigit():
                    messagebox.showerror("Erro de ID", "O ID deve ser um número inteiro.")
                    return
                id_para_deletar = id_da_entry

            elif selecao_treeview:
                # Caso 3: Somente a linha da Treeview está selecionada
                id_para_deletar = self.listaCli.item(selecao_treeview[0], 'values')[0]

            else:
                # Caso 4: Nenhum método de exclusão foi fornecido
                messagebox.showwarning("Erro de Seleção", 
                                    "Por favor, digite um ID ou selecione uma linha para exclusão.")
                return

            # Busca o nome do cliente no banco de dados usando o ID
            self.conecta_bd()
            self.cursor.execute("SELECT cli_nome FROM clientes WHERE cli_id = ?", (id_para_deletar,))
            resultado_busca = self.cursor.fetchone()
            
            if resultado_busca is None:
                messagebox.showerror("Erro", f"O cliente com o ID: {id_para_deletar} não foi encontrado.")
                return

            nome_cliente = resultado_busca[0]
            
            # Confirmação com o nome e o ID do cliente
            confirmacao = messagebox.askyesno("Confirmar Exclusão", 
                                            f"Tem certeza que deseja excluir o cliente?\n\nNome: {nome_cliente}\nID: {id_para_deletar}?")
            
            if not confirmacao:
                return

            # Executa a exclusão no banco de dados
            self.cursor.execute("DELETE FROM clientes WHERE cli_id = ?", (id_para_deletar,))
            self.conn.commit()
            
            messagebox.showinfo("Sucesso", "Cliente excluído com sucesso!")
            self.select_listaClientes(self.listaCli)

        except Exception as e:
            messagebox.showerror("Erro", f"Não foi possível excluir o cliente.\n\nErro: {e}")
            
        finally:
            self.desconecta_bd()
    
    def add_cliente(self):
        
        campos_obrigatorios = [
            self.nome_entry,
            self.nacionalidade_combo,
            self.estado_civil_combo,
            self.profissao_entry,
            self.rg_entry,
            self.cpf_entry,
            self.cep_entry,
            self.uf_entry,
            self.cidade_entry,
            self.rua_entry,
            self.n_rua_entry,
            self.bairro_entry
        ]

        campos_com_validacao = campos_obrigatorios + [self.email_entry] + [self.cnpj_reu_entry]

        for campo in campos_obrigatorios:
            # Primeiro, verifica se o campo está vazio
            if not campo.get().strip():
                messagebox.showwarning("Erro de Validação", "Por favor, preencha todos os campos obrigatórios (*).")
                return # Sai da função se encontrar um campo vazio
            
        # Verifica o estado 'invalid' de todos os campos que possuem validação    
        for campo in campos_com_validacao:
            # Em seguida, verifica o estado 'invalid' (para validações de formato)
            if 'invalid' in campo.state():
                messagebox.showwarning("Erro de Validação", "Por favor, preencha corretamente os campos marcados em vermelho.")
                return # Sai da função se houver um campo com formato inválido

        self.variaveis_cli()
        
        self.conecta_bd()

        try:
            self.cursor.execute("""
                INSERT INTO clientes (
                    cli_nome, cli_nacionalidade, cli_estado_civil, cli_profissao, cli_rg, 
                    cli_cpf, cli_cep, cli_uf, cli_cidade, cli_logradouro, cli_n_rua, 
                    cli_bairro, cli_telefone, cli_email, cli_nome_reu, cli_cnpj_reu
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                self.nome_completo, self.nacionalidade, self.estado_civil, self.profissao,self.rg,
                self.cpf, self.cep, self.uf, self.cidade, self.logradouro, self.n_rua, 
                self.bairro, self.telefone, self.email, self.nome_reu, self.cnpj_reu
            ))
            
            self.conn.commit()
            messagebox.showinfo("Sucesso", "Cliente adicionado com sucesso!")
            self.select_listaClientes(self.listaCli)
            self.limpa_cliente()
            
        except sqlite3.IntegrityError as e:
            messagebox.showerror("Erro", f"Cliente já pode existir no banco de dados, ou algum valor está fora do padrão.\n\nErro: {e}")
            
        finally:
            self.desconecta_bd()
   
    def busca_cliente(self, search_entry, lista_treeview, event=None):
        nome_busca = search_entry.get().strip()
        
        # Se o campo de busca estiver vazio, recarrega a lista completa
        if not nome_busca:
            self.select_listaClientes(lista_treeview)
            return

        try:
            self.conecta_bd()
            lista_treeview.delete(*lista_treeview.get_children())
            
            # Remove caracteres de formatação para checar se é um número (CPF)
            busca_formatada = nome_busca.replace('.', '').replace('-', '').replace('/', '')

            if busca_formatada.isdigit():
                # A busca é um número, então busca por CPF
                sql_query = """
                    SELECT * FROM clientes WHERE cli_cpf LIKE ? ORDER BY cli_nome ASC
                """
                search_param = (f"%{busca_formatada}%",)
                self.cursor.execute(sql_query, search_param)
                resultados_finais = self.cursor.fetchall()
                
                for i in resultados_finais:
                    lista_treeview.insert("", END, values=i)
                
            else:
                # A busca é um texto, então busca por nome (sem acentos)
                nome_normalizado_busca = str(unicodedata.normalize('NFKD', nome_busca).encode('ascii', 'ignore').decode('utf-8')).lower()
                sql_query = """
                    SELECT * FROM clientes ORDER BY cli_nome ASC
                """
                
                self.cursor.execute(sql_query)
                resultados_do_banco = self.cursor.fetchall()
                
                resultados_finais = []
                for cliente in resultados_do_banco:
                    nome_cliente_bd = cliente[1]
                    nome_bd_normalizado = str(unicodedata.normalize('NFKD', nome_cliente_bd).encode('ascii', 'ignore').decode('utf-8')).lower()
                    
                    if nome_normalizado_busca in nome_bd_normalizado:
                        resultados_finais.append(cliente)

                for i in resultados_finais:
                    lista_treeview.insert("", END, values=i)
                    
        except Exception as e:
            print(f"Ocorreu um erro durante a busca em tempo real: {e}")

        finally:
            self.desconecta_bd()
    
    def select_listaClientes(self, lista_treeview):
        lista_treeview.delete(*lista_treeview.get_children())
        self.conecta_bd()
        lista = self.cursor.execute(""" SELECT * FROM clientes ORDER BY cli_nome ASC""")

        for i in lista:
            lista_treeview.insert("", END, values=i)
        self.desconecta_bd()

    def busca_documento(self, search_entry, lista_treeview, event=None):
        nome_busca = search_entry.get().strip()

        # Se o campo de busca estiver vazio, recarrega a lista completa
        if not nome_busca:
            self.select_listaDocumentos(lista_treeview)
            return

        try:
            self.conecta_bd()
            lista_treeview.delete(*lista_treeview.get_children())
            
            # Checa se a entrada é um número (ID) ou texto (nome do documento)
            if nome_busca.isdigit():
                # A busca é um número, então busca por ID
                sql_query = """
                    SELECT doc_id, doc_nome, doc_tipo FROM documentos WHERE doc_id LIKE ? ORDER BY doc_nome ASC
                """
                search_param = (f"%{nome_busca}%",)
                self.cursor.execute(sql_query, search_param)
                resultados_finais = self.cursor.fetchall()
                
                for i in resultados_finais:
                    lista_treeview.insert("", END, values=i)
                
            else:
                # A busca é um texto, então busca por nome (sem acentos)
                nome_normalizado_busca = str(unicodedata.normalize('NFKD', nome_busca).encode('ascii', 'ignore').decode('utf-8')).lower()
                
                self.cursor.execute("SELECT doc_id, doc_nome, doc_tipo FROM documentos ORDER BY doc_nome ASC")
                resultados_do_banco = self.cursor.fetchall()
                
                resultados_finais = []
                for documento in resultados_do_banco:
                    nome_doc_bd = documento[1]
                    nome_doc_bd_normalizado = str(unicodedata.normalize('NFKD', nome_doc_bd).encode('ascii', 'ignore').decode('utf-8')).lower()
                    
                    if nome_normalizado_busca in nome_doc_bd_normalizado:
                        resultados_finais.append(documento)

                for i in resultados_finais:
                    lista_treeview.insert("", END, values=i)
                    
        except Exception as e:
            print(f"Ocorreu um erro durante a busca em tempo real: {e}")

        finally:
            self.desconecta_bd()
    
    def busca_cli_doc(self, event=None):
        nome_busca = self.cd_busca_entry.get().strip()

        # Se o campo de busca estiver vazio, recarrega a lista completa
        if not nome_busca:
            self.select_listaCliDoc()
            return

        try:
            self.conecta_bd()
            self.listaCliDoc.delete(*self.listaCliDoc.get_children())
            
            # Remove caracteres de formatação para checar se é um número (CPF)
            busca_formatada_cpf = re.sub(r'\D', '', nome_busca)

            # A busca é um número se o texto limpo não estiver vazio
            if busca_formatada_cpf:
                # A busca é por CPF.
                query = """
                    SELECT 
                        dg.dg_id, 
                        c.cli_nome, 
                        c.cli_cpf, 
                        dg.dg_nome, 
                        d.doc_tipo, 
                        dg.dg_data_criacao
                    FROM 
                        documento_gerado AS dg
                    INNER JOIN 
                        clientes AS c ON dg.fk_clientes_id = c.cli_id
                    INNER JOIN 
                        documentos AS d ON dg.fk_documentos_id = d.doc_id
                    WHERE
                        c.cli_cpf LIKE ?
                    ORDER BY 
                        dg.dg_data_criacao DESC
                """
                search_param = (f"%{busca_formatada_cpf}%",)
                self.cursor.execute(query, search_param)
                resultados_finais = self.cursor.fetchall()
                
                for documento in resultados_finais:
                    self.listaCliDoc.insert("", "end", values=documento)
                    
            else:
                # A busca é um texto. Buscamos todos os registros para filtrar em Python
                # e ignorar acentos, como você já fazia.
                query = """
                    SELECT 
                        dg.dg_id, 
                        c.cli_nome, 
                        c.cli_cpf, 
                        dg.dg_nome, 
                        d.doc_tipo, 
                        dg.dg_data_criacao
                    FROM 
                        documento_gerado AS dg
                    INNER JOIN 
                        clientes AS c ON dg.fk_clientes_id = c.cli_id
                    INNER JOIN 
                        documentos AS d ON dg.fk_documentos_id = d.doc_id
                """
                self.cursor.execute(query)
                resultados_do_banco = self.cursor.fetchall()
                
                resultados_finais = []
                nome_normalizado_busca = unicodedata.normalize('NFKD', nome_busca).encode('ascii', 'ignore').decode('utf-8').lower()
                
                for documento in resultados_do_banco:
                    nome_cliente_bd = documento[1]
                    nome_documento_bd = documento[3]
                    
                    nome_cli_bd_normalizado = unicodedata.normalize('NFKD', nome_cliente_bd).encode('ascii', 'ignore').decode('utf-8').lower()
                    nome_doc_bd_normalizado = unicodedata.normalize('NFKD', nome_documento_bd).encode('ascii', 'ignore').decode('utf-8').lower()
                    
                    # Verifica se a busca está em um dos dois campos
                    if nome_normalizado_busca in nome_cli_bd_normalizado or nome_normalizado_busca in nome_doc_bd_normalizado:
                        resultados_finais.append(documento)

                for item in resultados_finais:
                    self.listaCliDoc.insert("", "end", values=item)
                            
        except Exception as e:
            print(f"Ocorreu um erro durante a busca em tempo real: {e}")

        finally:
            self.desconecta_bd()
    
    def select_listaDocumentos(self, lista_treeview):
        lista_treeview.delete(*lista_treeview.get_children())
        self.conecta_bd()
        
        documentos = self.cursor.execute(""" 
            SELECT doc_id, doc_nome, doc_tipo FROM documentos ORDER BY doc_nome ASC
        """)

        for doc in documentos:
            lista_treeview.insert("", "end", values=doc)
            
        self.desconecta_bd()
    
    def select_listaCliDoc(self):
        self.listaCliDoc.delete(*self.listaCliDoc.get_children())
        self.conecta_bd()

        # Query SQL corrigida para buscar os dados de todas as colunas
        documentos = self.cursor.execute(""" 
            SELECT 
                dg.dg_id, 
                c.cli_nome, 
                c.cli_cpf, 
                dg.dg_nome, 
                d.doc_tipo, 
                dg.dg_data_criacao 
            FROM 
                documento_gerado AS dg
            INNER JOIN 
                clientes AS c ON dg.fk_clientes_id = c.cli_id
            INNER JOIN 
                documentos AS d ON dg.fk_documentos_id = d.doc_id
            ORDER BY 
                dg.dg_data_criacao DESC
        """)

        for doc in documentos:
            self.listaCliDoc.insert("", "end", values=doc)

        self.desconecta_bd()
    
    def del_cli_doc(self, event):
        selecoes_treeview = self.listaCliDoc.selection()

        # 1. Verifica se alguma linha está selecionada
        if not selecoes_treeview:
            messagebox.showwarning("Erro de Seleção", 
                                "Por favor, selecione uma ou mais linhas para exclusão.")
            return

        # 2. Pergunta de confirmação antes de iniciar o processo
        confirmacao = messagebox.askyesno("Confirmar Exclusão", 
                                        f"Tem certeza que deseja excluir os {len(selecoes_treeview)} documento(s) gerado(s)?")
        
        if not confirmacao:
            return

        try:
            self.conecta_bd()
            for item_id in selecoes_treeview:
                # Obtém o valor da primeira coluna (ID do documento)
                id_para_deletar = self.listaCliDoc.item(item_id, 'values')[0]

                # Executa a exclusão no banco de dados para cada ID
                self.cursor.execute("DELETE FROM documento_gerado WHERE dg_id = ?", (id_para_deletar,))
                
            self.conn.commit()
            
            messagebox.showinfo("Sucesso", f"{len(selecoes_treeview)} documento(s) excluído(s) com sucesso!")
            self.select_listaCliDoc()  # Recarrega a Treeview após a exclusão
            
        except Exception as e:
            self.conn.rollback()  # Desfaz as operações em caso de erro
            messagebox.showerror("Erro", f"Não foi possível excluir o(s) documento(s).\n\nErro: {e}")
            
        finally:
            self.desconecta_bd()
    
    def upt_documento(self):
        id_para_atualizar = self.id_doc_entry.get().strip()

        try:
            # Valida o ID
            if not id_para_atualizar:
                messagebox.showwarning("Erro de ID", "Por favor, digite o ID do documento para atualizar.")
                return
            
            if not id_para_atualizar.isdigit():
                messagebox.showerror("Erro de ID", "O ID deve ser um número inteiro.")
                return

            # Pega os novos valores dos campos de entrada
            nome_doc_novo = self.nome_doc_entry.get().strip()
            tipo_doc_novo = self.tipo_combo.get().strip()
            caminho_doc_novo = self.caminho_doc_entry.get().strip()

            self.conecta_bd()
            self.cursor.execute("SELECT * FROM documentos WHERE doc_id = ?", (id_para_atualizar,))
            dados_atuais = self.cursor.fetchone()

            if not dados_atuais:
                messagebox.showerror("Erro", f"Documento com o ID: {id_para_atualizar} não foi encontrado no banco de dados.")
                return

            nome_doc_atual = dados_atuais[1]
            tipo_doc_atual = dados_atuais[2]
            conteudo_arquivo_atual = dados_atuais[3]

            nome_doc_final = nome_doc_novo if nome_doc_novo else nome_doc_atual
            tipo_doc_final = tipo_doc_novo if tipo_doc_novo else tipo_doc_atual
            
            conteudo_arquivo_final = conteudo_arquivo_atual
            
            if caminho_doc_novo:
                if not os.path.exists(caminho_doc_novo):
                    messagebox.showerror("Erro no Arquivo", f"O novo arquivo '{caminho_doc_novo}' não foi encontrado.")
                    return
                with open(caminho_doc_novo, 'rb') as arquivo:
                    conteudo_arquivo_final = arquivo.read()

                # --- Bloco de Verificação de Arquivo Duplicado ---
                # Busca por arquivos existentes, EXCLUINDO o documento atual
                self.cursor.execute("SELECT doc_arquivo FROM documentos WHERE doc_id != ?", (id_para_atualizar,))
                arquivos_existentes = self.cursor.fetchall()

                for arquivo_existente in arquivos_existentes:
                    if conteudo_arquivo_final == arquivo_existente[0]:
                        messagebox.showerror("Erro de Duplicidade", "Este arquivo já existe no banco de dados em outro documento.")
                        return
                # --- Fim do Bloco de Verificação ---
            
            # --- Bloco de Confirmação ---
            confirmacao = messagebox.askyesno("Confirmar Atualização", 
                                            f"Tem certeza que deseja atualizar o documento?\n"
                                            f"\nNome: {nome_doc_atual}\nID: {id_para_atualizar}")
            
            if not confirmacao:
                return
            # --- Fim do Bloco de Confirmação ---
            
            # Executa a atualização no banco de dados
            self.cursor.execute("""
                UPDATE documentos SET 
                    doc_nome = ?, doc_tipo = ?, doc_arquivo = ? 
                WHERE doc_id = ?
            """, (
                nome_doc_final, tipo_doc_final, conteudo_arquivo_final, id_para_atualizar
            ))
            
            self.conn.commit()
            messagebox.showinfo("Sucesso", "Documento atualizado com sucesso!")
            self.select_listaDocumentos(self.listaDoc)
            self.limpa_documento()
            
        except sqlite3.IntegrityError as e:
            messagebox.showerror("Erro de Integridade", "Documento com esse nome já existe no banco de dados.")
        except Exception as e:
            messagebox.showerror("Erro na Atualização", f"Não foi possível atualizar o documento.\n\nErro: {e}")
            
        finally:
            self.desconecta_bd()
    
    def add_documento(self):
        campos_obrigatorios = [
            self.nome_doc_entry,
            self.tipo_combo,
            self.caminho_doc_entry
        ]

        for campo in campos_obrigatorios:
            if not campo.get().strip():
                messagebox.showwarning("Erro de Validação", "Por favor, preencha todos os campos obrigatórios (*).")
                return
                
        try:
            nome_doc = self.nome_doc_entry.get().strip()
            tipo_doc = self.tipo_combo.get().strip()
            caminho_doc = self.caminho_doc_entry.get().strip()
            
            # Verifica se o arquivo existe no caminho fornecido
            if not os.path.exists(caminho_doc):
                messagebox.showerror("Erro no Arquivo", f"O arquivo '{caminho_doc}' não foi encontrado.")
                return

            # Abre e lê o arquivo em modo binário
            with open(caminho_doc, 'rb') as arquivo:
                conteudo_novo_arquivo = arquivo.read()

            self.conecta_bd()

            # 1. Busca todos os conteúdos de arquivos já existentes no banco de dados
            self.cursor.execute("SELECT doc_arquivo FROM documentos")
            arquivos_existentes = self.cursor.fetchall()
            
            # 2. Compara o novo arquivo com os existentes
            for arquivo_existente in arquivos_existentes:
                if conteudo_novo_arquivo == arquivo_existente[0]:
                    messagebox.showerror("Erro de Duplicidade", "Este arquivo já existe no banco de dados.")
                    return

            # Se não encontrou duplicatas, prossegue com a inserção
            self.cursor.execute("""
                INSERT INTO documentos (
                    doc_nome, doc_tipo, doc_arquivo
                )
                VALUES (?, ?, ?)
            """, (
                nome_doc, tipo_doc, conteudo_novo_arquivo
            ))
            
            self.conn.commit()
            messagebox.showinfo("Sucesso", "Documento adicionado com sucesso!")
            self.select_listaDocumentos(self.listaDoc)
            self.limpa_documento()
                
        except sqlite3.IntegrityError as e:
            messagebox.showerror("Erro de Integridade", "Documento com esse nome já existe no banco de dados.")
                
        except Exception as e:
            messagebox.showerror("Erro", f"Não foi possível adicionar o documento.\n\nErro: {e}")
                
        finally:
            self.desconecta_bd()

    def del_documento(self, event):
        id_para_deletar = None

        # 1. Tenta obter o ID do campo id_entry (se estiver preenchido)
        id_da_entry = self.id_doc_entry.get().strip()

        # 2. Tenta obter o ID da linha selecionada na Treeview
        selecao_treeview = self.listaDoc.selection() # Assumindo que a Treeview de documentos se chama self.listaDoc

        try:
            if id_da_entry and selecao_treeview:
                # Caso 1: Ambos estão preenchidos, verifica se são iguais
                id_da_linha = self.listaDoc.item(selecao_treeview[0], 'values')[0]
                if str(id_da_entry) != str(id_da_linha):
                    messagebox.showwarning("Aviso de Conflito", 
                                            "O ID no campo (digitado) é diferente do ID da linha selecionada.\n"
                                            "Por favor, use apenas um dos métodos para exclusão.")
                    return
                else:
                    id_para_deletar = id_da_entry

            elif id_da_entry:
                # Caso 2: Somente o campo id_entry está preenchido
                if not id_da_entry.isdigit():
                    messagebox.showerror("Erro de ID", "O ID deve ser um número inteiro.")
                    return
                id_para_deletar = id_da_entry

            elif selecao_treeview:
                # Caso 3: Somente a linha da Treeview está selecionada
                id_para_deletar = self.listaDoc.item(selecao_treeview[0], 'values')[0]

            else:
                # Caso 4: Nenhum método de exclusão foi fornecido
                messagebox.showwarning("Erro de Seleção", 
                                        "Por favor, digite um ID ou selecione uma linha para exclusão.")
                return

            # Busca o nome do documento no banco de dados usando o ID
            self.conecta_bd()
            self.cursor.execute("SELECT doc_nome FROM documentos WHERE doc_id = ?", (id_para_deletar,))
            resultado_busca = self.cursor.fetchone()
            
            if resultado_busca is None:
                messagebox.showerror("Erro", f"O documento com o ID: {id_para_deletar} não foi encontrado.")
                return

            nome_documento = resultado_busca[0]
            
            # Confirmação com o nome e o ID do documento
            confirmacao = messagebox.askyesno("Confirmar Exclusão", 
                                                f"Tem certeza que deseja excluir o documento?\n\nNome: {nome_documento}\nID: {id_para_deletar}?")
            
            if not confirmacao:
                return

            # Executa a exclusão no banco de dados
            self.cursor.execute("DELETE FROM documentos WHERE doc_id = ?", (id_para_deletar,))
            self.conn.commit()
            
            messagebox.showinfo("Sucesso", "Documento excluído com sucesso!")
            self.select_listaDocumentos(self.listaDoc) # Chama a função que recarrega a Treeview de documentos

        except Exception as e:
            messagebox.showerror("Erro", f"Não foi possível excluir o documento.\n\nErro: {e}")
            
        finally:
            self.desconecta_bd()

    def escolher_arquivo(self):
        #Abre o explorer
        caminho = filedialog.askopenfilename(
            title="Selecione um documento",
            filetypes=[("Documentos Word", "*.docx"), ("Todos os arquivos", "*.*")]
        )
        if caminho:
            self.caminho_doc_entry.delete(0, END)
            self.caminho_doc_entry.insert(0, caminho)
    
    def exportar(self, file_type, treeview_widget, tabela_nome, colunas_tabela):
        selecao_treeview = treeview_widget.selection()

        try:
            if not selecao_treeview:
                messagebox.showwarning("Erro de Seleção", 
                                    "Por favor, selecione uma linha para exportação.")
                return

            id_para_exportar = treeview_widget.item(selecao_treeview[0], 'values')[0]
            nome_coluna, arquivo_coluna, id_coluna = colunas_tabela

            self.conecta_bd()
            
            # Constrói a consulta SQL de forma dinâmica e segura
            sql_query = f"SELECT {nome_coluna}, {arquivo_coluna} FROM {tabela_nome} WHERE {id_coluna} = ?"
            
            self.cursor.execute(sql_query, (id_para_exportar,))
            resultado = self.cursor.fetchone()
            
            if resultado is None:
                messagebox.showerror("Erro", f"O documento com o ID: {id_para_exportar} não foi encontrado na tabela '{tabela_nome}'.")
                return
                
            nome_doc_bd, conteudo_doc_bd = resultado

            if file_type == 'word':
                extensao = ".docx"
                filetypes = [("Arquivos do Word", "*.docx"), ("Todos os arquivos", "*.*")]
                caminho_salvar = filedialog.asksaveasfilename(
                    defaultextension=extensao,
                    initialfile=f"{nome_doc_bd.replace(' ', '_').replace('.', '')}{extensao}",
                    title=f"Salvar Arquivo",
                    filetypes=filetypes
                )
                
                if not caminho_salvar: return
                    
                with open(caminho_salvar, 'wb') as arquivo:
                    arquivo.write(conteudo_doc_bd)
            
            elif file_type == 'pdf':
                extensao = ".pdf"
                filetypes = [("Arquivos PDF", "*.pdf"), ("Todos os arquivos", "*.*")]
                
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
                    temp_docx.write(conteudo_doc_bd)
                    caminho_temp_docx = temp_docx.name
                
                caminho_salvar = filedialog.asksaveasfilename(
                    defaultextension=extensao,
                    initialfile=f"{nome_doc_bd.replace(' ', '_').replace('.', '')}{extensao}",
                    title=f"Salvar Arquivo",
                    filetypes=filetypes
                )
                
                if not caminho_salvar:
                    os.remove(caminho_temp_docx)
                    return

                # --- Indicador de Progresso ---
                largura_tela = self.root.winfo_screenwidth()
                altura_tela = self.root.winfo_screenheight()
                largura_janela = 300
                altura_janela = 50
                posicao_x = (largura_tela // 2) - (largura_janela // 2)
                posicao_y = (altura_tela // 2) - (altura_janela // 2)

                janela_progresso = Toplevel(self.root)
                janela_progresso.title("Convertendo...")
                janela_progresso.geometry(f"{largura_janela}x{altura_janela}+{posicao_x}+{posicao_y}")
                janela_progresso.transient(self.root)
                janela_progresso.grab_set()

                label_progresso = Label(janela_progresso, text=f"Aguarde, convertendo {nome_doc_bd} para PDF...")
                label_progresso.pack(pady=15)
                janela_progresso.update_idletasks()
                
                try:
                    convert(caminho_temp_docx, caminho_salvar)
                    time.sleep(1)
                    
                except Exception as e:
                    messagebox.showerror("Erro de Conversão", f"Não foi possível converter o arquivo para PDF.\nVerifique se o Microsoft Word (Windows) ou LibreOffice (Linux) está instalado.\n\nErro: {e}")
                    return
                
                finally:
                    if os.path.exists(caminho_temp_docx):
                        os.remove(caminho_temp_docx)
                    janela_progresso.destroy()
                    
            else:
                messagebox.showerror("Erro", "Tipo de arquivo não suportado. Use 'word' ou 'pdf'.")
                return

            messagebox.showinfo("Sucesso", f"O documento '{nome_doc_bd}{extensao}' foi exportado com sucesso!")

        except Exception as e:
            messagebox.showerror("Erro", f"Não foi possível exportar o documento.\n\nErro: {e}")
            
        finally:
            self.desconecta_bd()
    
    def formatar_cpf(self, cpf):
        """Aplica a máscara de CPF."""
        cpf_limpo = ''.join(filter(str.isdigit, str(cpf)))
        if len(cpf_limpo) == 11:
            return f"{cpf_limpo[:3]}.{cpf_limpo[3:6]}.{cpf_limpo[6:9]}-{cpf_limpo[9:]}"
        return cpf_limpo

    def formatar_rg(self, rg):
        """Aplica a máscara de RG."""
        rg_limpo = ''.join(filter(str.isdigit, str(rg)))
        if len(rg_limpo) == 9:
            return f"{rg_limpo[:2]}.{rg_limpo[2:5]}.{rg_limpo[5:8]}-{rg_limpo[8:]}"
        return rg_limpo

    def formatar_cep(self, cep):
        """Aplica a máscara de CEP."""
        cep_limpo = ''.join(filter(str.isdigit, str(cep)))
        if len(cep_limpo) == 8:
            return f"{cep_limpo[:5]}-{cep_limpo[5:]}"
        return cep_limpo

    def formatar_telefone(self, tel):
        """Aplica a máscara de telefone (com ou sem DDD)."""
        tel_limpo = ''.join(filter(str.isdigit, str(tel)))
        if len(tel_limpo) == 11: # Ex: (99) 99999-9999
            return f"({tel_limpo[:2]}) {tel_limpo[2:7]}-{tel_limpo[7:]}"
        elif len(tel_limpo) == 10: # Ex: (99) 9999-9999
            return f"({tel_limpo[:2]}) {tel_limpo[2:6]}-{tel_limpo[6:]}"
        return tel_limpo

    def formatar_cnpj(self, cnpj):
        """Aplica a máscara de CNPJ."""
        cnpj_limpo = ''.join(filter(str.isdigit, str(cnpj)))
        if len(cnpj_limpo) == 14:
            return f"{cnpj_limpo[:2]}.{cnpj_limpo[2:5]}.{cnpj_limpo[5:8]}/{cnpj_limpo[8:12]}-{cnpj_limpo[12:]}"
        return cnpj_limpo

    def gerar_documento_cli_doc(self, lista_clientes, lista_documentos, entry_nome):
        selecao_cliente = lista_clientes.selection()
        selecao_documento = lista_documentos.selection()
        nome_arquivo_gerado = entry_nome.get().strip()

        if not selecao_cliente:
            messagebox.showwarning("Aviso", "Por favor, selecione um cliente.")
            return
        if not selecao_documento:
            messagebox.showwarning("Aviso", "Por favor, selecione um documento.")
            return
        if not nome_arquivo_gerado:
            messagebox.showwarning("Aviso", "Por favor, digite um nome para o arquivo gerado.")
            return

        cli_id = lista_clientes.item(selecao_cliente[0], 'values')[0]
        doc_id = lista_documentos.item(selecao_documento[0], 'values')[0]

        self.conecta_bd()
        try:
            # 1. Resgata os dados do cliente e o arquivo do documento
            self.cursor.execute("SELECT * FROM clientes WHERE cli_id = ?", (cli_id,))
            dados_cliente = self.cursor.fetchone()
            
            self.cursor.execute("SELECT doc_arquivo FROM documentos WHERE doc_id = ?", (doc_id,))
            arquivo_doc_binario = self.cursor.fetchone()[0]

            if not dados_cliente or not arquivo_doc_binario:
                messagebox.showerror("Erro", "Cliente ou documento não encontrado no banco de dados.")
                return

            # 2. Configura o idioma para o mês em português
            try:
                locale.setlocale(locale.LC_TIME, 'pt_BR.utf8')
            except locale.Error:
                locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil.1252')

            # Mapeamento dos dados do cliente para as variáveis do documento (valores em str)
            mapa_variaveis = {
                "{{ID}}": str(dados_cliente[0]),
                "{{NOME}}": str(dados_cliente[1] or ""),
                "{{NACIONALIDADE}}": str(dados_cliente[2] or ""),
                "{{ESTADO_CIVIL}}": str(dados_cliente[3] or ""),
                "{{PROFISSAO}}": str(dados_cliente[4] or ""),
                "{{RG}}": str(self.formatar_rg(dados_cliente[5]) if dados_cliente[5] else ""),
                "{{CPF}}": str(self.formatar_cpf(dados_cliente[6]) if dados_cliente[6] else ""),
                "{{CEP}}": str(self.formatar_cep(dados_cliente[7]) if dados_cliente[7] else ""),
                "{{UF}}": str(dados_cliente[8] or ""),
                "{{CIDADE}}": str(dados_cliente[9] or ""),
                "{{LOGRADOURO}}": str(dados_cliente[10] or ""),
                "{{N_RUA}}": str(dados_cliente[11] or ""),
                "{{BAIRRO}}": str(dados_cliente[12] or ""),
                "{{TELEFONE}}": str(self.formatar_telefone(dados_cliente[13]) if dados_cliente[13] else ""),
                "{{EMAIL}}": str(dados_cliente[14] or ""),
                "{{NOME_REU}}": str(dados_cliente[15] or ""),
                "{{CNPJ_REU}}": str(self.formatar_cnpj(dados_cliente[16]) if dados_cliente[16] else ""),
                "{{DIA}}": str(datetime.now().day),
                "{{MES}}": datetime.now().strftime('%B').capitalize(),
                "{{ANO}}": str(datetime.now().year)
            }

            # 3. Gera um arquivo temporário para trabalhar com o documento
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                temp_file.write(arquivo_doc_binario)
                caminho_temp = temp_file.name

            doc_modelo = Document(caminho_temp)

            # ------------ função robusta de substituição que preserva estilo ------------
            def substituir_variaveis_no_paragrafo(paragrafo, variaveis):
                runs = list(paragrafo.runs)  # snapshot
                if not runs:
                    return

                textos = [r.text for r in runs]
                combinado = "".join(textos)
                if not any(k in combinado for k in variaveis):
                    return  # nada a fazer

                # boundaries: posições cumulativas dos runs
                boundaries = [0]
                for t in textos:
                    boundaries.append(boundaries[-1] + len(t))

                def find_run_index(char_idx):
                    # retorna índice do run que contém o char_idx
                    return bisect_right(boundaries, char_idx) - 1

                def snapshot_run(run):
                    f = run.font
                    snap = {
                        "name": f.name,
                        "size": f.size,
                        "bold": f.bold,
                        "italic": f.italic,
                        "underline": f.underline,
                        "color": None
                    }
                    try:
                        snap["color"] = f.color.rgb
                    except Exception:
                        snap["color"] = None
                    return snap

                def apply_snapshot_to_run(run, snap):
                    f = run.font
                    # só aplica quando o atributo não é None, para não sobrescrever heranças de estilo
                    if snap.get("name") is not None:
                        f.name = snap["name"]
                    if snap.get("size") is not None:
                        f.size = snap["size"]
                    if snap.get("bold") is not None:
                        f.bold = snap["bold"]
                    if snap.get("italic") is not None:
                        f.italic = snap["italic"]
                    if snap.get("underline") is not None:
                        f.underline = snap["underline"]
                    if snap.get("color") is not None:
                        try:
                            f.color.rgb = snap["color"]
                        except Exception:
                            pass

                segments = []  # lista de (texto, estilo_snap)

                i = 0
                L = len(combinado)
                # busca próxima ocorrência de qualquer placeholder
                while i < L:
                    next_pos = None
                    next_key = None
                    for key in variaveis:
                        pos = combinado.find(key, i)
                        if pos != -1 and (next_pos is None or pos < next_pos):
                            next_pos = pos
                            next_key = key
                    if next_pos is None:
                        # não há mais placeholders: pega o resto mantendo estilos dos runs
                        start = i
                        end = L
                        # percorre runs cobrindo [start,end)
                        start_run = find_run_index(start)
                        end_run = find_run_index(end - 1) if end - 1 >= 0 else start_run
                        for r_idx in range(start_run, end_run + 1):
                            r = runs[r_idx]
                            r_start = boundaries[r_idx]
                            slice_start = max(start, r_start) - r_start
                            slice_end = min(end, boundaries[r_idx + 1]) - r_start
                            texto_slice = r.text[slice_start:slice_end]
                            if texto_slice:
                                segments.append((texto_slice, snapshot_run(r)))
                        break
                    if next_pos > i:
                        # adiciona o texto normal entre i e next_pos (respeitando runs)
                        start = i
                        end = next_pos
                        start_run = find_run_index(start)
                        end_run = find_run_index(end - 1)
                        for r_idx in range(start_run, end_run + 1):
                            r = runs[r_idx]
                            r_start = boundaries[r_idx]
                            slice_start = max(start, r_start) - r_start
                            slice_end = min(end, boundaries[r_idx + 1]) - r_start
                            texto_slice = r.text[slice_start:slice_end]
                            if texto_slice:
                                segments.append((texto_slice, snapshot_run(r)))
                        i = next_pos
                    else:
                        # placeholder começa em i
                        # escolhe estilo do run onde o placeholder começa
                        start_run_idx = find_run_index(i)
                        estilo_sub = snapshot_run(runs[start_run_idx])
                        valor = str(variaveis[next_key])
                        segments.append((valor, estilo_sub))
                        i += len(next_key)

                # reconstruir os runs do parágrafo com os segmentos
                # remove runs existentes
                for r in list(paragrafo.runs):
                    r._element.getparent().remove(r._element)

                # adiciona novos runs de acordo com segmentos (aplicando estilo)
                for texto, estilo in segments:
                    novo_run = paragrafo.add_run(texto)
                    apply_snapshot_to_run(novo_run, estilo)
            # -------------------------------------------------------------------------

            # Aplica substituições em todos os parágrafos do documento
            for paragrafo in doc_modelo.paragraphs:
                substituir_variaveis_no_paragrafo(paragrafo, mapa_variaveis)

            # Aplica em tabelas
            for tabela in doc_modelo.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for paragrafo in celula.paragraphs:
                            substituir_variaveis_no_paragrafo(paragrafo, mapa_variaveis)

            # 5. Salva o documento modificado em um novo arquivo temporário
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as novo_doc_temp:
                doc_modelo.save(novo_doc_temp.name)
                novo_doc_temp.seek(0)
                novo_arquivo_binario = novo_doc_temp.read()
                caminho_novo_doc_temp = novo_doc_temp.name

            # 6. Armazena o novo documento no banco de dados
            data_criacao = datetime.now().strftime("%d-%m-%Y")
            
            self.cursor.execute(
                "INSERT INTO documento_gerado (fk_clientes_id, fk_documentos_id, dg_nome, dg_data_criacao, dg_arquivo) VALUES (?, ?, ?, ?, ?)",
                (cli_id, doc_id, nome_arquivo_gerado, data_criacao, novo_arquivo_binario)
            )
            self.conn.commit()
            
            messagebox.showinfo("Sucesso", f"Documento '{nome_arquivo_gerado}' gerado e salvo com sucesso!")

            # 7. Abre o arquivo gerado automaticamente (Windows, macOS, Linux)
            try:
                if os.name == 'nt':
                    if self.abrir_apos_gerar_var.get():
                        os.startfile(caminho_novo_doc_temp)
                elif sys.platform == 'darwin':
                    subprocess.call(['open', caminho_novo_doc_temp])
                else:
                    subprocess.call(['xdg-open', caminho_novo_doc_temp])
            except Exception:
                # falhar ao abrir não é crítico; apenas continue
                pass

        except Exception as e:
            messagebox.showerror("Erro", f"Não foi possível gerar o documento.\nErro: {e}")
            
        finally:
            # Limpa o arquivo temporário original (mantemos o novo para que o usuário possa abrir)
            if 'caminho_temp' in locals() and os.path.exists(caminho_temp):
                try:
                    os.remove(caminho_temp)
                except Exception:
                    pass
            
            self.desconecta_bd()

    
    def criar_treeview_generica(self, pai, colunas, cabecalhos, larguras_colunas, relx, rely, relwidth, relheight):
        # Cria a Treeview
        treeview = ttk.Treeview(pai, height=3, columns=colunas, show="headings")

        # Configura os cabeçalhos e larguras das colunas
        for i, (cabecalho, largura) in enumerate(zip(cabecalhos, larguras_colunas)):
            treeview.heading(f"#{i+1}", text=cabecalho)
            treeview.column(f"#{i+1}", width=largura)

        # Posiciona a Treeview
        treeview.place(relx=relx, rely=rely, relwidth=relwidth, relheight=relheight)

        # Cria e configura a barra de rolagem vertical
        scroll_y = ttk.Scrollbar(pai, orient='vertical', command=treeview.yview)
        treeview.configure(yscrollcommand=scroll_y.set)
        scroll_y.place(relx=relx + relwidth, rely=rely, relwidth=0.02, relheight=relheight)

        # Cria e configura a barra de rolagem horizontal
        scroll_x = ttk.Scrollbar(pai, orient='horizontal', command=treeview.xview)
        treeview.configure(xscrollcommand=scroll_x.set)
        scroll_x.place(relx=relx, rely=rely + relheight, relwidth=relwidth, relheight=0.02)

        # Lógica de scroll com o mouse
        def _on_mousewheel_vertical(event):
            treeview.yview_scroll(-3 * int(event.delta / 2), "units")

        def _on_mousewheel_horizontal(event):
            treeview.xview_scroll(-1 * int(event.delta / 2), "units")

        treeview.bind("<MouseWheel>", _on_mousewheel_vertical)
        treeview.bind("<Shift-MouseWheel>", _on_mousewheel_horizontal)
        
        return treeview

    def lista_documentos(self):
        #Declaração dos parâmetros
        colunas = ("col1", "col2", "col3")
        cabecalhos = ("ID", "Nome", "Tipo")
        larguras = (50, 350, 100)
        relx, rely, relwidth, relheight = 0.02, 0.3, 0.95, 0.6
        
        self.listaDoc = self.criar_treeview_generica(self.frame_documento, colunas, cabecalhos, larguras, relx, rely, relwidth, relheight)

        self.listaDoc.bind("<Double-1>", self.OnDoubleClick_Doc)

    def lista_clientes(self):
        #Declaração dos parâmetros
        colunas = ("col1","col2","col3","col4","col5","col6","col7","col8","col9","col10","col11","col12","col13","col14","col15","col16","col17")
        cabecalhos = ("ID", "Nome Completo", "Nacionalidade", "Estado Civil", "Profissão", "RG", "CPF", "CEP", "UF", "Cidade", "Rua", "Nº Rua", "Bairro", "Telefone", "Email", "Nome Completo Réu", "CNPJ Réu")
        larguras = (35, 250, 100, 110, 200, 100, 110, 85, 40, 150, 220, 150, 200, 110, 230, 400, 130)
        relx, rely, relwidth, relheight = 0.02, 0.51, 0.95, 0.46
        
        self.listaCli = self.criar_treeview_generica(self.frame_cliente, colunas, cabecalhos, larguras, relx, rely, relwidth, relheight)

        self.listaCli.bind("<Double-1>", self.OnDoubleClick_Cli)

    def lista_gerar_cli(self):
        #Declaração dos parâmetros
        colunas = ("col1","col2","col3","col4","col5","col6","col7","col8","col9","col10","col11","col12","col13","col14","col15","col16","col17")
        cabecalhos = ("ID", "Nome Completo", "Nacionalidade", "Estado Civil", "Profissão", "RG", "CPF", "CEP", "UF", "Cidade", "Rua", "Nº Rua", "Bairro", "Telefone", "Email", "Nome Completo Réu", "CNPJ Réu")
        larguras = (35, 250, 100, 110, 200, 100, 110, 85, 40, 150, 220, 150, 200, 110, 230, 400, 130)
        relx, rely, relwidth, relheight = 0.02, 0.11, 0.95, 0.3
        
        self.listaGerarCli = self.criar_treeview_generica(self.frame_gerar_doc, colunas, cabecalhos, larguras, relx, rely, relwidth, relheight)

    def lista_gerar_doc(self):
        #Declaração dos parâmetros
        colunas = ("col1", "col2", "col3")
        cabecalhos = ("ID", "Nome", "Tipo")
        larguras = (50, 350, 100)
        relx, rely, relwidth, relheight = 0.02, 0.56, 0.95, 0.3
        
        self.listaGerarDoc = self.criar_treeview_generica(self.frame_gerar_doc, colunas, cabecalhos, larguras, relx, rely, relwidth, relheight)
    
    def lista_clientes_doc(self):
        #Declaração dos parâmetros
        colunas = ("col1","col2","col3", "col4", "col5", "col6")
        cabecalhos = ("ID", "Nome Cliente", "CPF Cliente", "Nome Documento", "Tipo Documento", "Data Gerada")
        larguras = (25, 150, 100, 150, 75, 75)
        relx, rely, relwidth, relheight = 0.02, 0.17, 0.95, 0.73
        
        self.listaCliDoc = self.criar_treeview_generica(self.frame_clientes_doc, colunas, cabecalhos, larguras, relx, rely, relwidth, relheight)
    
    def gerenciar_clientes(self):
        #Configuração de estilo para as Entrys
        self.style = ttk.Style()
        self.style.configure("Big.TEntry", font=("Helvetica", 20))

        #Label e Entry do id
        self.lb_id = ttk.Label(self.frame_cliente, text = "ID")
        self.lb_id.place(relx= 0.04, rely=0.02)
        self.id_entry = ttk.Entry(self.frame_cliente, style="Big.TEntry")
        self.id_entry.place(relx= 0.03, rely=0.05, relwidth=0.05)

        #Label e Entry do nome
        self.lb_nome = ttk.Label(self.frame_cliente, text = "Nome Completo*")
        self.lb_nome.place(relx= 0.12, rely=0.02)
        self.nome_entry = ttk.Entry(self.frame_cliente, style="Big.TEntry")
        self.nome_entry.place(relx= 0.11, rely=0.05, relwidth=0.30)

        #Label e Combobox da nacionalidade
        self.lb_nacionalidade = ttk.Label(self.frame_cliente, text = "Nacionalidade*")
        self.lb_nacionalidade.place(relx= 0.45, rely=0.02)
        opc_nacionalidade = ["afegão", "afegã", "albanês", "albanesa", "alemão", "alemã", "americano", "americana", "andorrano", "andorrana", "angolano", "angolana", "antiguano", "antiguana", "argentino", "argentina", "armênio", "armênia", "australiano", "australiana", "austríaco", "austríaca", "azerbaijano", "azerbaijana", "bahamense", "bangladeshiano", "bangladeshiana", "barbadiano", "barbadiana", "belga", "belizenho", "belizenha", "beninense", "bielorrusso", "bielorrussa", "boliviano", "boliviana", "bósnio", "bósnia", "botsuanês", "botsuanesa", "brasileiro", "brasileira", "britânico", "britânica", "búlgaro", "búlgara", "burquinense", "burundês", "burundesa", "butanês", "butanesa", "cabo-verdiano", "cabo-verdiana", "camaronês", "camaronesa", "cambojano", "cambojana", "canadense", "catariano", "catarina", "chileno", "chilena", "chinês", "chinesa", "cingapuriano", "cingapuriana", "colombiano", "colombiana", "congolês", "congolesa", "coreano do norte", "coreana do norte", "coreano do sul", "coreana do sul", "costarriquenho", "costarriquenha", "croata", "cubano", "cubana", "dinamarquês", "dinamarquesa", "dominicano", "dominicana", "egípcio", "egípcia", "equatoriano", "equatoriana", "eritreu", "eritreia", "escocês", "escocesa", "eslovaco", "eslovaca", "esloveno", "eslovena", "espanhol", "espanhola", "estoniano", "estoniana", "etíope", "filipino", "filipina", "finlandês", "finlandesa", "francês", "francesa", "gabonês", "gabonesa", "galês", "galesa", "ganês", "ganesa", "georgiano", "georgiana", "grego", "grega", "guatemalteco", "guatemalteca", "guianês", "guianesa", "guineense", "haitiano", "haitiana", "holandês", "holandesa", "hondurenho", "hondurenha", "húngaro", "húngara", "iemenita", "indiano", "indiana", "indonésio", "indonésia", "inglês", "inglesa", "iraquiano", "iraquiana", "iraniano", "iraniana", "irlandês", "irlandesa", "islandês", "islandesa", "israelense", "italiano", "italiana", "jamaicano", "jamaicana", "japonês", "japonesa", "jordano", "jordana", "kazakhstanês", "kazakhstanesa", "keniano", "keniana", "kiribati", "kuwaitiano", "kuwaitiana", "letão", "letona", "libanês", "libanesa", "liberiano", "liberiana", "líbio", "líbia", "liechtensteiniano", "liechtensteiniana", "lituano", "lituana", "luxemburguês", "luxemburguesa", "macedônio", "macedônia", "malaio", "malaia", "malawiano", "malawiana", "maliano", "maliana", "maltês", "maltesa", "marroquino", "marroquina", "mauriciano", "mauriciana", "mexicano", "mexicana", "moçambicano", "moçambicana", "moldávio", "moldávia", "monegasco", "monegasca", "mongol", "montenegrino", "montenegrina", "namibiano", "namibiana", "nepalês", "nepalesa", "nicaraguense", "nigeriano", "nigeriana", "norueguês", "norueguesa", "neozelandês", "neozelandesa", "omanês", "omanesa", "paquistanês", "paquistanesa", "palestino", "palestina", "panamenho", "panamenha", "papua nova guiné", "paraguaio", "paraguaia", "peruano", "peruana", "polonês", "polonesa", "portorriquenho", "portorriquenha", "português", "portuguesa", "qatari", "qatari", "queniano", "queniana", "quirguiz", "quirguiz", "romeno", "romena", "ruandês", "ruandesa", "russo", "russa", "salvadorenho", "salvadorenha", "samoano", "samoana", "sanmarinense", "sanmarinense", "saudita", "saudita", "senegalês", "senegalesa", "sérvio", "sérvia", "somaliano", "somaliana", "sudanês", "sudanesa", "sueco", "sueca", "suíço", "suíça", "surinamês", "surinamesa", "tailandês", "tailandesa", "tanzaniano", "tanzaniana", "timorense", "timorense", "togolês", "togolesa", "turco", "turca", "turcomano", "turcomana", "ucraniano", "ucraniana", "ugandês", "ugandesa", "uruguaio", "uruguaia", "uzbeque", "uzbeque", "venezuelano", "venezuelana", "vietnamita", "vietnamita", "zambiano", "zambiana", "zimbabuano", "zimbabuana"]
        self.nacionalidade_combo = AutocompleteEntry(opc_nacionalidade, self.frame_cliente, style="Big.TEntry", listboxLength=6, width=23, matchesFunction=self.matches)
        self.nacionalidade_combo.place(relx=0.44, rely=0.05, relwidth=0.13)

        #Label e Combobox do estado civil
        self.lb_estado_civil = ttk.Label(self.frame_cliente, text = "Estado Civil*")
        self.lb_estado_civil.place(relx= 0.61, rely=0.02)
        opc_estado_civil = ["solteiro", "solteira", "casado", "casada", "divorciado", "divorciada", "viúvo", "viúva", "separado", "separada", "união estável", "separação judicial"]
        self.estado_civil_combo = AutocompleteEntry(opc_estado_civil, self.frame_cliente, style="Big.TEntry", listboxLength=6, width=23, matchesFunction=self.matches)
        self.estado_civil_combo.place(relx=0.6, rely=0.05, relwidth=0.12)

        #Label e Entry da profissão
        self.lb_profissao = ttk.Label(self.frame_cliente, text = "Profissão*")
        self.lb_profissao.place(relx= 0.76, rely=0.02)
        self.profissao_entry = ttk.Entry(self.frame_cliente, style="Big.TEntry")
        self.profissao_entry.place(relx= 0.75, rely=0.05, relwidth=0.22)

        #Label e Entry do RG
        self.lb_rg = ttk.Label(self.frame_cliente, text = "RG*")
        self.lb_rg.place(relx= 0.04, rely=0.11)
        self.rg_entry = ttk.Entry(self.frame_cliente, style="Big.TEntry")
        self.rg_entry.place(relx= 0.03, rely=0.14, relwidth=0.15)
        self.rg_entry.bind("<KeyRelease>", self.valida_rg)

        #Label e Entry do CPF
        self.lb_cpf = ttk.Label(self.frame_cliente, text = "CPF*")
        self.lb_cpf.place(relx= 0.22, rely=0.11)
        self.cpf_entry = ttk.Entry(self.frame_cliente, style="Big.TEntry")
        self.cpf_entry.place(relx= 0.21, rely=0.14, relwidth=0.15)
        self.cpf_entry.bind("<KeyRelease>", self.valida_cpf)

        #Label e Entry do CEP
        self.lb_cep = ttk.Label(self.frame_cliente, text = "CEP*")
        self.lb_cep.place(relx= 0.40, rely=0.11)
        self.cep_entry = ttk.Entry(self.frame_cliente, style="Big.TEntry")
        self.cep_entry.place(relx= 0.39, rely=0.14, relwidth=0.15)
        self.cep_entry.bind("<KeyRelease>", self.cepCorreios)

        #Label e Entry da UF
        self.lb_uf = ttk.Label(self.frame_cliente, text = "UF*")
        self.lb_uf.place(relx= 0.58, rely=0.11)
        self.uf_entry = ttk.Entry(self.frame_cliente, style="Big.TEntry")
        self.uf_entry.place(relx= 0.57, rely=0.14, relwidth=0.05)

        #Label e Entry da cidade
        self.lb_cidade = ttk.Label(self.frame_cliente, text = "Cidade*")
        self.lb_cidade.place(relx= 0.66, rely=0.11)
        self.cidade_entry = ttk.Entry(self.frame_cliente, style="Big.TEntry")
        self.cidade_entry.place(relx= 0.65, rely=0.14, relwidth=0.32)

        #Label e Entry da rua
        self.lb_rua = ttk.Label(self.frame_cliente, text = "Rua*")
        self.lb_rua.place(relx= 0.04, rely=0.20)
        self.rua_entry = ttk.Entry(self.frame_cliente, style="Big.TEntry")
        self.rua_entry.place(relx= 0.03, rely=0.23, relwidth=0.31)

        #Label e Entry do numero da rua
        self.lb_n_rua = ttk.Label(self.frame_cliente, text = "Nº Rua*")
        self.lb_n_rua.place(relx= 0.38, rely=0.20)
        self.n_rua_entry = ttk.Entry(self.frame_cliente, style="Big.TEntry")
        self.n_rua_entry.place(relx= 0.37, rely=0.23, relwidth=0.1)

        #Label e Entry do bairro
        self.lb_bairro = ttk.Label(self.frame_cliente, text = "Bairro*")
        self.lb_bairro.place(relx= 0.51, rely=0.20)
        self.bairro_entry = ttk.Entry(self.frame_cliente, style="Big.TEntry")
        self.bairro_entry.place(relx= 0.5, rely=0.23, relwidth=0.31)

        #Label e Entry do telefone
        self.lb_telefone = ttk.Label(self.frame_cliente, text = "Telefone")
        self.lb_telefone.place(relx= 0.85, rely=0.20)
        self.telefone_entry = ttk.Entry(self.frame_cliente, style="Big.TEntry")
        self.telefone_entry.place(relx= 0.84, rely=0.23, relwidth=0.13)

        #Label e Entry do email
        self.lb_email = ttk.Label(self.frame_cliente, text = "Email")
        self.lb_email.place(relx= 0.04, rely=0.29)
        self.email_entry = ttk.Entry(self.frame_cliente, style="Big.TEntry")
        self.email_entry.place(relx= 0.03, rely=0.32, relwidth=0.25)
        self.email_entry.bind("<KeyRelease>", self.valida_email)

        #Label e Entry do nome do réu
        self.lb_nome_reu = ttk.Label(self.frame_cliente, text = "Nome Completo Réu")
        self.lb_nome_reu.place(relx= 0.32, rely=0.29)
        self.nome_reu_entry = ttk.Entry(self.frame_cliente, style="Big.TEntry")
        self.nome_reu_entry.place(relx= 0.31, rely=0.32, relwidth=0.4)

        #Label e Entry do CNPJ ou CPF do réu
        self.lb_cnpj_reu = ttk.Label(self.frame_cliente, text = "CNPJ/CPF Réu")
        self.lb_cnpj_reu.place(relx= 0.75, rely=0.29)
        self.cnpj_reu_entry = ttk.Entry(self.frame_cliente, style="Big.TEntry")
        self.cnpj_reu_entry.place(relx= 0.74, rely=0.32, relwidth=0.23)
        self.cnpj_reu_entry.bind("<KeyRelease>", self.valida_cpf_cnpj)

        #Botão adicionar cliente
        self.bt_add_cliente = ttk.Button(self.frame_cliente, text="Adicionar Cliente", style='Accent.TButton', command=self.add_cliente)
        self.bt_add_cliente.place(relx= 0.03, rely=0.44, relwidth=0.17)

        #Botão update cliente
        self.bt_update_cliente = ttk.Button(self.frame_cliente, text="Atualizar Cliente", style='Accent.TButton', command=self.upt_cliente)
        self.bt_update_cliente.place(relx= 0.23, rely=0.44, relwidth=0.17)

        #Botão deletar cliente
        self.bt_del_cliente = ttk.Button(self.frame_cliente, text="Deletar Cliente", style='Accent.TButton', command=lambda: self.del_cliente(None))
        self.bt_del_cliente.place(relx= 0.43, rely=0.44, relwidth=0.17)

        #Label e Entry da pesquisa por nome ou cpf do cliente
        self.lb_nome_busca = ttk.Label(self.frame_cliente, text = "Buscar nome/CPF")
        self.lb_nome_busca.place(relx= 0.67, rely=0.41)
        self.nome_busca_entry = ttk.Entry(self.frame_cliente, style="Big.TEntry")
        self.nome_busca_entry.place(relx= 0.66, rely=0.44, relwidth=0.31)
        self.nome_busca_entry.bind("<KeyRelease>", lambda event: self.busca_cliente(self.nome_busca_entry ,self.listaCli, event))

        #Botão limpar campos
        self.bt_limpar = ttk.Button(self.frame_cliente, text="Limpar", command=self.limpa_cliente)
        self.bt_limpar.place(relx= 0.84, rely=0.38, relwidth=0.11)

        #Criação da Treeview
        self.lista_clientes()
        self.select_listaClientes(self.listaCli)
        self.setup_enter_bindings("frame_cliente")
        self.listaCli.bind("<Delete>", self.del_cliente)
        self.listaCli.bind("<BackSpace>", self.del_cliente)
    
    def gerenciar_documentos(self):
        #Configuração de estilo para as Entrys
        self.style = ttk.Style()
        self.style.configure("Big.TEntry", font=("Helvetica", 20))

        #Label e Entry do ID do documento
        self.lb_id_doc = ttk.Label(self.frame_documento, text="ID")
        self.lb_id_doc.place(relx= 0.04, rely=0.02)
        self.id_doc_entry = ttk.Entry(self.frame_documento, style="Big.TEntry")
        self.id_doc_entry.place(relx= 0.03, rely=0.05, relwidth=0.05)
        
        #Label e Entry do nome do documento
        self.lb_nome_doc = ttk.Label(self.frame_documento, text="Nome do Documento*")
        self.lb_nome_doc.place(relx= 0.12, rely=0.02)
        self.nome_doc_entry = ttk.Entry(self.frame_documento, style="Big.TEntry")
        self.nome_doc_entry.place(relx= 0.11, rely=0.05, relwidth=0.4)

        #Label e Combobox do tipo do documento
        self.lb_tipo = ttk.Label(self.frame_documento, text = "Tipo*")
        self.lb_tipo.place(relx= 0.55, rely=0.02)
        opc_tipo = ["Contrato", "Declaração", "Petição", "Tese", "Procuração"]
        self.tipo_combo = AutocompleteEntry(opc_tipo, self.frame_documento, style="Big.TEntry", listboxLength=6, width=23, matchesFunction=self.matches)
        self.tipo_combo.place(relx=0.54, rely=0.05, relwidth=0.25)

        #Label e Entry do caminho do documento
        self.lb_texto_caminho_doc = ttk.Label(self.frame_documento, text="Caminho do documento*: ")
        self.lb_texto_caminho_doc.place(relx= 0.04, rely=0.12)
        self.caminho_doc_entry = ttk.Entry(self.frame_documento, style="Big.TEntry")
        self.caminho_doc_entry.place(relx= 0.22, rely=0.11, relwidth=0.57)
        #Botão para abrir o explorer
        self.btn_escolher_doc = ttk.Button(self.frame_documento, text="Escolher Arquivo", command=self.escolher_arquivo)
        self.btn_escolher_doc.place(relx= 0.8, rely=0.11, relwidth=0.17)

        #Botão adicionar documento
        self.bt_add_documento = ttk.Button(self.frame_documento, text="Adicionar Documento", style='Accent.TButton', command=self.add_documento)
        self.bt_add_documento.place(relx= 0.03, rely=0.23, relwidth=0.17)

        #Botão atualizar documento
        self.bt_upt_documento = ttk.Button(self.frame_documento, text="Atualizar Documento", style='Accent.TButton', command=self.upt_documento)
        self.bt_upt_documento.place(relx= 0.23, rely=0.23, relwidth=0.17)
        
        #Botão deletar documento
        self.bt_del_documento = ttk.Button(self.frame_documento, text="Deletar Documento", style='Accent.TButton', command=lambda: self.del_documento(None))
        self.bt_del_documento.place(relx= 0.43, rely=0.23, relwidth=0.17)

        #Label e Entry da pesquisa por nome do documento
        self.lb_nome_doc_busca = ttk.Label(self.frame_documento, text="Buscar por nome")
        self.lb_nome_doc_busca.place(relx= 0.66, rely=0.2)
        self.nome_doc_busca_entry = ttk.Entry(self.frame_documento, style="Big.TEntry")
        self.nome_doc_busca_entry.place(relx= 0.65, rely=0.23, relwidth=0.33)
        self.nome_doc_busca_entry.bind("<KeyRelease>", lambda event: self.busca_documento(self.nome_doc_busca_entry, self.listaDoc, event))

        #Botão limpar campos documento
        self.bt_limpar_documento = ttk.Button(self.frame_documento, text="Limpar", command=self.limpa_documento)
        self.bt_limpar_documento.place(relx= 0.82, rely=0.17, relwidth=0.11)

        metadados_doc_modelo = ("doc_nome", "doc_arquivo", "doc_id")
        
        #Botão exportar para Word
        self.bt_export_word = ttk.Button(self.frame_documento, text="Exportar Word", style='Accent.TButton', command=lambda: self.exportar("word", self.listaDoc, "documentos", metadados_doc_modelo))
        self.bt_export_word.place(relx= 0.60, rely=0.94, relwidth=0.17)

        #Botão exportar para PDF
        self.bt_export_pdf = ttk.Button(self.frame_documento, text="Exportar PDF", style='Accent.TButton', command=lambda: self.exportar("pdf", self.listaDoc, "documentos", metadados_doc_modelo))
        self.bt_export_pdf.place(relx= 0.80, rely=0.94, relwidth=0.17)
        
        #Criação da Treeview
        self.lista_documentos()
        self.select_listaDocumentos(self.listaDoc)
        self.setup_enter_bindings("frame_documento")
        self.listaDoc.bind("<Delete>", self.del_documento)
        self.listaDoc.bind("<BackSpace>", self.del_documento)
    
    def clientes_doc(self):
        #Configuração de estilo para as Entrys
        self.style = ttk.Style()
        self.style.configure("Big.TEntry", font=("Helvetica", 20))

        #Label e Entry da pesquisa por Cliente
        self.lb_cd_busca = ttk.Label(self.frame_clientes_doc, text="Busca nome/CPF cliente")
        self.lb_cd_busca.place(relx= 0.03, rely=0.02)
        self.cd_busca_entry = ttk.Entry(self.frame_clientes_doc, style="Big.TEntry")
        self.cd_busca_entry.place(relx= 0.02, rely=0.05, relwidth=0.4)
        self.cd_busca_entry.bind("<KeyRelease>", self.busca_cli_doc)

        #Botão deletar registro
        self.bt_del_cd = ttk.Button(self.frame_clientes_doc, text="Deletar Registro", style='Accent.TButton', command=lambda: self.del_cli_doc(None))
        self.bt_del_cd.place(relx= 0.02, rely=0.11, relwidth=0.17)
        
        #Criação da Treeview
        self.lista_clientes_doc()
        self.select_listaCliDoc()
        self.listaCliDoc.bind("<Delete>", self.del_cli_doc)
        self.listaCliDoc.bind("<BackSpace>", self.del_cli_doc)

        metadados_doc_gerado = ("dg_nome", "dg_arquivo", "dg_id")
        
        #Botão exportar para Word
        self.bt_export_word_cd = ttk.Button(self.frame_clientes_doc, text="Exportar Word", style='Accent.TButton', command=lambda: self.exportar("word", self.listaCliDoc, "documento_gerado", metadados_doc_gerado))
        self.bt_export_word_cd.place(relx= 0.60, rely=0.94, relwidth=0.17)

        #Botão exportar para PDF
        self.bt_export_pdf_cd = ttk.Button(self.frame_clientes_doc, text="Exportar PDF", style='Accent.TButton', command=lambda: self.exportar("pdf", self.listaCliDoc, "documento_gerado", metadados_doc_gerado))
        self.bt_export_pdf_cd.place(relx= 0.80, rely=0.94, relwidth=0.17)

    def gerar_doc(self):        
        #Configuração de estilo para as Entrys
        self.style = ttk.Style()
        self.style.configure("Big.TEntry", font=("Helvetica", 20))

        #Label e Entry da pesquisa por Cliente
        self.lb_gerar_cli_busca = ttk.Label(self.frame_gerar_doc, text="Busca nome/CPF cliente")
        self.lb_gerar_cli_busca.place(relx= 0.03, rely=0.02)
        self.gerar_cli_busca_entry = ttk.Entry(self.frame_gerar_doc, style="Big.TEntry")
        self.gerar_cli_busca_entry.place(relx= 0.02, rely=0.05, relwidth=0.4)
        self.gerar_cli_busca_entry.bind("<KeyRelease>", lambda event: self.busca_cliente(self.gerar_cli_busca_entry, self.listaGerarCli, event))

        self.lista_gerar_cli()
        self.select_listaClientes(self.listaGerarCli)

        #Label e Entry da pesquisa por Documento
        self.lb_gerar_doc_busca = ttk.Label(self.frame_gerar_doc, text="Busca nome Documento")
        self.lb_gerar_doc_busca.place(relx= 0.03, rely=0.47)
        self.gerar_doc_busca_entry = ttk.Entry(self.frame_gerar_doc, style="Big.TEntry")
        self.gerar_doc_busca_entry.place(relx= 0.02, rely=0.5, relwidth=0.4)
        self.gerar_doc_busca_entry.bind("<KeyRelease>", lambda event: self.busca_documento(self.gerar_doc_busca_entry, self.listaGerarDoc, event))

        self.lista_gerar_doc()
        self.select_listaDocumentos(self.listaGerarDoc)

        #Label e Entry do nome do Documento Gerado
        self.lb_gerar_nome = ttk.Label(self.frame_gerar_doc, text="Nome do documento*")
        self.lb_gerar_nome.place(relx= 0.03, rely=0.89)
        self.gerar_nome_entry = ttk.Entry(self.frame_gerar_doc, style="Big.TEntry")
        self.gerar_nome_entry.place(relx= 0.02, rely=0.92, relwidth=0.5)
        
        #Botão gerar Word
        self.bt_gerar_word = ttk.Button(self.frame_gerar_doc, text="Gerar Word", style='Accent.TButton', command=lambda: self.gerar_documento_cli_doc(self.listaGerarCli, self.listaGerarDoc, self.gerar_nome_entry))
        self.bt_gerar_word.place(relx= 0.80, rely=0.92, relwidth=0.17)

        #
        self.abrir_apos_gerar_var = tk.BooleanVar(value=True)
        self.cb_abrir_apos_gerar = ttk.Checkbutton(self.frame_gerar_doc, text="Abrir documento após gerar", variable=self.abrir_apos_gerar_var, style="TCheckbutton")
        self.cb_abrir_apos_gerar.place(relx=0.55, rely=0.92)
        
class App(Funcs):
    def __init__(self):
        #Configurações de inicialização
        self.root = root
        self.tela()
        self.frames_da_tela()
        self.widgets_frame1()
        self.montaTabelas()
        root.mainloop()

    def tela(self):
        #Configurações da tela
        self.root.title("Sistema Advogados")
        sv_ttk.set_theme("light") 
        self.root.geometry("1280x1024")
        self.root.resizable(True, True)
        self.root.maxsize(width=1440, height=900)
        self.root.minsize(width=1024, height=768)
    
    def frames_da_tela(self):
        #Frame do menu de opções
        self.frame_menu = ttk.Frame(self.root, style='Card.TFrame')
        self.frame_menu.place(relx= 0.02, rely=0.02, relwidth=0.2, relheight=0.96)

        #Frame da tela Gerenciar Clientes
        self.frame_cliente = ttk.Frame(self.root, style='Card.TFrame')
        self.frame_cliente.place(relx= 0.24, rely=0.02, relwidth=0.74, relheight=0.96)

        #Frame da tela Gerenciar Documentos
        self.frame_documento = ttk.Frame(self.root, style='Card.TFrame')
        self.frame_documento.place(relx= 0.24, rely=0.02, relwidth=0.74, relheight=0.96)

        #Frame da tela Clientes/Documentos
        self.frame_clientes_doc = ttk.Frame(self.root, style='Card.TFrame')
        self.frame_clientes_doc.place(relx= 0.24, rely=0.02, relwidth=0.74, relheight=0.96)

        #Frame da tela Gerar Contratos e Afins
        self.frame_gerar_doc = ttk.Frame(self.root, style='Card.TFrame')
        self.frame_gerar_doc.place(relx= 0.24, rely=0.02, relwidth=0.74, relheight=0.96)

    def widgets_frame1(self):
        #Logo da minha empresa
        self.img = tk.PhotoImage(file="logo.png")
        self.img_logo = Label(self.frame_menu, image=self.img)
        self.img_logo.place(relx= 0.1, rely=0.02, relwidth=0.8, relheight=0.33)

        #Botão Gerenciar Clientes
        self.bt_clientes = ttk.Button(self.frame_menu, text="Gerenciar Clientes", command=lambda: [self.gerenciar_clientes(), self.frame_cliente.lift(), self.root.after(1, self.root.focus_set)])
        self.bt_clientes.place(relx= 0.05, rely=0.37, relwidth=0.9, relheight=0.05)

        #Botão Gerenciar Documentos
        self.bt_documentos = ttk.Button(self.frame_menu, text="Gerenciar Documentos", command=lambda: [self.gerenciar_documentos(), self.frame_documento.lift(), self.root.after(1, self.root.focus_set)])
        self.bt_documentos.place(relx= 0.05, rely=0.44, relwidth=0.9, relheight=0.05)

        #Botão Clientes/Documentos
        self.bt_clientes_doc = ttk.Button(self.frame_menu, text="Clientes/Documentos", command=lambda: [self.clientes_doc(), self.frame_clientes_doc.lift(), self.root.after(1, self.root.focus_set)])
        self.bt_clientes_doc.place(relx= 0.05, rely=0.51, relwidth=0.9, relheight=0.05)

        #Botão Gerar Contratos e Afins
        self.bt_contratos = ttk.Button(self.frame_menu, text="Gerar Contratos e Afins", command=lambda: [self.gerar_doc(), self.frame_gerar_doc.lift(), self.root.after(1, self.root.focus_set)])
        self.bt_contratos.place(relx= 0.05, rely=0.58, relwidth=0.9, relheight=0.05)

        #Botão Gerar Petição
        self.bt_peticao = ttk.Button(self.frame_menu, text="Gerar Petição")
        self.bt_peticao.place(relx= 0.05, rely=0.65, relwidth=0.9, relheight=0.05)

        self.switch = ttk.Checkbutton(self.frame_menu, text="Dark theme", style="Switch.TCheckbutton", command=sv_ttk.toggle_theme)
        self.switch.place(relx= 0.05, rely=0.8)

        #Botão Sair
        self.bt_sair = ttk.Button(self.frame_menu, text="Sair", command=self.root.destroy, style='Accent.TButton')
        self.bt_sair.place(relx= 0.05, rely=0.93, relwidth=0.9, relheight=0.05)

App()