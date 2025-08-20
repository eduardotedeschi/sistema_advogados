from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy
import os

def create_page_break_element():
    """Cria um elemento de parágrafo contendo apenas uma quebra de página."""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p

def remove_leading_page_breaks(body):
    """Remove parágrafos iniciais que sejam apenas quebras de página."""
    while len(body) > 0 and is_paragraph_page_break(body[0]):
        body.remove(body[0])

def is_paragraph_page_break(p):
    """Verifica se um elemento é apenas uma quebra de página."""
    if p.tag != qn('w:p'):
        return False
    for el in p.iter():
        if el.tag == qn('w:t') and el.text and el.text.strip():
            return False
    for el in p.iter():
        if el.tag == qn('w:br') and el.get(qn('w:type')) == 'page':
            return True
    return False

def juntar_docs(pasta, nome_arquivo_final):
    arquivos = [f for f in os.listdir(pasta) if f.lower().endswith('.docx')]
    if not arquivos:
        print("Nenhum arquivo .docx encontrado na pasta:", pasta)
        return

    documentos_ordenados = sorted(arquivos)
    documento_final = Document()

    # Remove parágrafo inicial vazio
    if documento_final.paragraphs:
        p = documento_final.paragraphs[0]._element
        p.getparent().remove(p)

    for i, arquivo in enumerate(documentos_ordenados):
        caminho_arquivo = os.path.join(pasta, arquivo)
        print("Processando:", caminho_arquivo)

        try:
            doc = Document(caminho_arquivo)

            if i > 0:
                # Adiciona quebra real no XML antes do próximo documento
                documento_final.element.body.append(create_page_break_element())

            for elemento in doc.element.body:
                documento_final.element.body.append(deepcopy(elemento))

        except Exception as e:
            print(f"Erro ao abrir {arquivo}: {e}")

    # Remove quebra de página inicial se sobrou alguma
    remove_leading_page_breaks(documento_final.element.body)

    caminho_saida = os.path.join(pasta, nome_arquivo_final)
    documento_final.save(caminho_saida)
    print(f'Documentos combinados com sucesso em: {caminho_saida}')


if __name__ == "__main__":
    caminho_pasta = r"C:\Users\dudu\Documents\CLIENTES WEB\daniel_adv\Projeto\banco_de_documentos"
    arquivo_saida = "documento_final.docx"
    juntar_docs(caminho_pasta, arquivo_saida)
