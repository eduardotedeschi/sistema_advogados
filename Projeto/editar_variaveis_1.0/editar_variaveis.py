from docx import Document
from datetime import datetime
import locale

# Define o idioma do sistema para português do Brasil
locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')

# ----------------- Inputs -----------------
nome = input("Nome completo do cliente: ").title() \
                              .replace(" Da ", " da ") \
                              .replace(" De ", " de ") \
                              .replace(" Do ", " do ") \
                              .replace(" Dos ", " dos ") \
                              .replace(" Das ", " das ") \
                              .replace(" E ", " e ") \
                              .replace(" Em ", " em ")
nacionalidade = input("Nacionalidade do cliente: ").lower()
estadoCivil = input("Estado civíl do cliente: ").lower()
profissao = input("Profissão do cliente: ")
rg = input("RG do cliente: ")
cpf = input("CPF do cliente: ")
logradouro = input("Rua do cliente: ").capitalize()
numeroResidencia = input("Número da residência do cliente: ")
bairro = input("Bairro do cliente: ").capitalize()
cidade = input("Cidade de cliente: ").title() \
                                     .replace(" Da ", " da ") \
                                     .replace(" De ", " de ") \
                                     .replace(" Do ", " do ") \
                                     .replace(" Dos ", " dos ") \
                                     .replace(" Das ", " das ") \
                                     .replace(" Em ", " em ")
uf = input("UF do cliente: ").upper()
cep = input("CEP do cliente: ")
telefone = input("Telefone do cliente: ")
email = input("Email do cliente: ").lower()
nomeReu = input("Nome completo do Réu: ").title() \
                              .replace(" Da ", " da ") \
                              .replace(" De ", " de ") \
                              .replace(" Do ", " do ") \
                              .replace(" Dos ", " dos ") \
                              .replace(" Das ", " das ") \
                              .replace(" E ", " e ") \
                              .replace(" Em ", " em ")
cnpjReu = input("CNPJ do Réu: ")

# ----------------- Documento -----------------
doc = Document(r"C:\Users\dudu\Documents\CLIENTES WEB\sistema_advogados\Projeto\editar_variaveis_1.0\banco_de_documentos\Contrato.docx")

# ----------------- Dados -----------------
dados = {
    "{{NOME}}": nome,
    "{{NACIONALIDADE}}": nacionalidade,
    "{{ESTADO_CIVIL}}": estadoCivil,
    "{{PROFISSAO}}": profissao,
    "{{RG}}": rg,
    "{{CPF}}": cpf,
    "{{LOGRADOURO}}": logradouro,
    "{{N_RUA}}": numeroResidencia,
    "{{BAIRRO}}": bairro,
    "{{CIDADE}}": cidade,
    "{{UF}}": uf,
    "{{CEP}}": cep,
    "{{TELEFONE}}": telefone,
    "{{EMAIL}}": email,
    "{{DIA}}": str(datetime.now().day),
    "{{MES}}": datetime.now().strftime('%B').capitalize(),
    "{{ANO}}": str(datetime.now().year),
    "{{NOMER_REU}}": nomeReu,
    "{{CNPJ_REU}}": cnpjReu
}

# ----------------- Função para substituir placeholders -----------------
def substituir_placeholders_paragrafo_robusto(paragraph, dados):
    """
    Substitui placeholders em um parágrafo mesmo se estiverem divididos em vários runs,
    preservando a formatação original.
    """
    # Cria lista de (run_text, run) para reconstruir
    runs = [(run.text, run) for run in paragraph.runs]
    full_text = ''.join(text for text, _ in runs)
    
    # Substitui todos os placeholders
    for chave, valor in dados.items():
        full_text = full_text.replace(chave, valor)
    
    # Limpa todos os runs
    for run in paragraph.runs:
        run.text = ''
    
    # Reconstrói o parágrafo distribuindo o texto pelos runs originais
    cursor = 0
    for text, run in runs:
        run_len = len(text)
        run.text = full_text[cursor:cursor + run_len]
        cursor += run_len
        if cursor >= len(full_text):
            break
    # Se sobrar texto (ex: placeholder maior que run original), adiciona ao último run
    if cursor < len(full_text):
        paragraph.runs[-1].text += full_text[cursor:]

def substituir_placeholders_elemento(elemento, dados):
    # Parágrafos
    if hasattr(elemento, "paragraphs"):
        for p in elemento.paragraphs:
            substituir_placeholders_paragrafo_robusto(p, dados)
    # Tabelas
    if hasattr(elemento, "tables"):
        for table in elemento.tables:
            for row in table.rows:
                for cell in row.cells:
                    substituir_placeholders_elemento(cell, dados)  # recursão

# ----------------- Substituir -----------------
substituir_placeholders_elemento(doc, dados)

# ----------------- Salvar -----------------
doc.save(rf"C:\Users\dudu\Documents\CLIENTES WEB\sistema_advogados\Projeto\editar_variaveis_1.0\banco_de_documentos\Contrato - {nome}.docx")

print(f"Documento gerado com sucesso: Contrato - {nome}.docx")
